from __future__ import annotations

from pathlib import Path
import os
import traceback
import uuid
import shutil
import json
import math
import re
import hashlib
import zipfile
from datetime import datetime
from typing import Any, Dict, List

from dotenv import load_dotenv
from flask import Flask, jsonify, render_template, request, abort, send_from_directory
from werkzeug.utils import secure_filename

from core.data_loader import KnowledgeBase, CaseRecord, case_sort_key
from core.risk_engine import generate_traceable_report
from core.llm_client import ask_llm, test_llm_connection
from core.case_parser import parse_case_file

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
DATA_PATH = Path(os.getenv("DATA_PATH", "data/knowledge_base.xlsx"))
if not DATA_PATH.is_absolute():
    DATA_PATH = BASE_DIR / DATA_PATH
# User-fed data is stored outside the project folder by default.
# This prevents losing added cases/articles when a new front-end version is unzipped.
# Override with USCC_DATA_DIR if you want a custom persistent path.
PERSISTENT_DATA_DIR = Path(os.getenv("USCC_DATA_DIR", Path.home() / ".uscc_scc_flask_data"))
PERSISTENT_DATA_DIR.mkdir(parents=True, exist_ok=True)
PROJECT_DATA_DIR = BASE_DIR / "data"
UPLOAD_DIR = PERSISTENT_DATA_DIR / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
CANDIDATE_BATCH_DIR = PERSISTENT_DATA_DIR / "candidate_batches"
CANDIDATE_BATCH_DIR.mkdir(parents=True, exist_ok=True)
USER_CASES_PATH = PERSISTENT_DATA_DIR / "user_cases.json"
DELETED_CASES_PATH = PERSISTENT_DATA_DIR / "deleted_cases.json"
ARTICLES_PATH = PERSISTENT_DATA_DIR / "articles.json"
CASE_TAGS_PATH = PERSISTENT_DATA_DIR / "case_tags.json"
MIGRATION_STATE_PATH = PERSISTENT_DATA_DIR / "migration_state.json"
MASCOT_SETTINGS_PATH = PERSISTENT_DATA_DIR / "mascot_settings.json"
API_CONFIG_PATH = PERSISTENT_DATA_DIR / "api_config.json"
API_CONFIG_HISTORY_PATH = PERSISTENT_DATA_DIR / "api_config_history.json"
LIBRARY_STATE_PATH = PERSISTENT_DATA_DIR / "library_state.json"
KB_DIGEST_PATH = PERSISTENT_DATA_DIR / "knowledge_digest.json"
PERSISTENCE_REPAIR_LOG_PATH = PERSISTENT_DATA_DIR / "persistence_repair_log.json"
ARTICLES_DELETED_MARKER = PERSISTENT_DATA_DIR / "articles_deleted_all.marker"
ARTICLES_V34_CLEARED_MARKER = PERSISTENT_DATA_DIR / "articles_v34_cleared.marker"

app = Flask(__name__)
app.config["SECRET_KEY"] = os.getenv("FLASK_SECRET_KEY", "dev-only-secret")
app.config["MAX_CONTENT_LENGTH"] = int(os.getenv("MAX_UPLOAD_MB", "512")) * 1024 * 1024

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff", ".dcm"}
CASE_TABLE_EXTENSIONS = {".xlsx", ".xls", ".csv", ".docx", ".doc", ".pdf", ".txt", ".md"}
ARTICLE_EXTENSIONS = {".txt", ".docx", ".pdf", ".xlsx", ".xls", ".csv", ".md"}
ALL_UPLOAD_EXTENSIONS = IMAGE_EXTENSIONS | CASE_TABLE_EXTENSIONS
MAX_BATCH_FILES = int(os.getenv("MAX_BATCH_FILES", "500"))
MAX_TEXT_CHARS_PER_FILE = int(os.getenv("MAX_TEXT_CHARS_PER_FILE", "500000"))

kb = KnowledgeBase(DATA_PATH)


def _record_to_saved_dict(rec: CaseRecord) -> Dict[str, Any]:
    d = _json_safe(rec.as_public_dict())
    d.update({
        "surgery": getattr(rec, "surgery", ""),
        "other_treatment": getattr(rec, "other_treatment", ""),
        "source_row": getattr(rec, "source_row", 0),
        "medical_images": _normalize_medical_images(getattr(rec, "medical_images", []) or []),
        "case_signature": getattr(rec, "case_signature", "") or "",
        "tags": _get_case_tags(rec),
    })
    return d



def _json_safe(value: Any) -> Any:
    """Return values that browser JSON.parse can always consume.
    Pandas/Excel data can contain NaN; some browsers reject JSON containing NaN.
    """
    if value is None:
        return ""
    if isinstance(value, float) and math.isnan(value):
        return ""
    if isinstance(value, dict):
        return {str(k): _json_safe(v) for k, v in value.items()}
    if isinstance(value, list):
        return [_json_safe(v) for v in value]
    return value


def _case_to_public_dict(rec: CaseRecord) -> Dict[str, Any]:
    data = rec.as_public_dict()
    data["medical_images"] = _normalize_medical_images(data.get("medical_images") or [])
    data["tags"] = _get_case_tags(rec)
    data["tag_text"] = "，".join(data["tags"])
    return _json_safe(data)


def _normalize_medical_images(images: Any) -> List[Dict[str, Any]]:
    """Keep case image references portable across app restarts."""
    normalized: List[Dict[str, Any]] = []
    for item in images if isinstance(images, list) else []:
        if not isinstance(item, dict):
            continue
        stored_as = Path(str(item.get("stored_as") or "")).name
        filename = str(item.get("filename") or stored_as or "医学影像")
        url = f"/uploads/{stored_as}" if stored_as else str(item.get("url") or "")
        normalized.append({
            **item,
            "type": "image",
            "filename": filename,
            "stored_as": stored_as,
            "url": url,
        })
    return normalized


def _normalize_source_pages(pages: Any) -> List[int]:
    normalized: List[int] = []
    for value in pages if isinstance(pages, list) else []:
        try:
            page = int(value)
        except (TypeError, ValueError):
            continue
        if page > 0 and page not in normalized:
            normalized.append(page)
    return sorted(normalized)


def _optional_int(value: Any) -> int | None:
    try:
        return int(value) if str(value or "").strip() else None
    except (TypeError, ValueError):
        return None


def _case_matches_query(rec: CaseRecord, query: str) -> bool:
    q = (query or "").strip().lower()
    if not q:
        return True
    public = _case_to_public_dict(rec)
    haystack = "\n".join([str(v) for v in public.values() if v is not None]).lower()
    # Support multiple keywords separated by spaces; Chinese continuous keywords still work.
    terms = [x for x in q.replace("，", " ").replace(",", " ").split() if x]
    if not terms:
        return q in haystack
    return all(term in haystack for term in terms)

def _load_deleted_case_ids() -> set[str]:
    rows = _read_json_with_backup(DELETED_CASES_PATH, [])
    return {str(x) for x in rows if str(x).strip()} if isinstance(rows, list) else set()


def _save_deleted_case_ids(ids: set[str]) -> None:
    _atomic_write_json(DELETED_CASES_PATH, sorted(ids))
    _save_library_snapshot()




def _load_case_tags() -> Dict[str, List[str]]:
    raw = _read_json_with_backup(CASE_TAGS_PATH, {})
    if isinstance(raw, dict):
        out: Dict[str, List[str]] = {}
        for cid, tags in raw.items():
            if isinstance(tags, str):
                items = [x.strip() for x in re.split(r"[,，;；\s]+", tags) if x.strip()]
            elif isinstance(tags, list):
                items = [str(x).strip() for x in tags if str(x).strip()]
            else:
                items = []
            out[str(cid)] = list(dict.fromkeys(items))[:12]
        return out
    return {}


def _save_case_tags(rows: Dict[str, List[str]]) -> None:
    clean = {str(cid): list(dict.fromkeys([str(t).strip() for t in tags if str(t).strip()]))[:12] for cid, tags in rows.items()}
    _atomic_write_json(CASE_TAGS_PATH, clean)


def _default_case_tags(rec: CaseRecord) -> List[str]:
    tags: List[str] = []
    for item in [getattr(rec, "sheet", ""), getattr(rec, "diagnosis", ""), getattr(rec, "ls", ""), getattr(rec, "tnm", "")]:
        text = str(item or "").strip()
        if not text:
            continue
        if len(text) > 18:
            continue
        if text not in tags:
            tags.append(text)
    return tags[:6]


def _get_case_tags(rec: CaseRecord) -> List[str]:
    """Return the single doctor-editable label used for knowledge-base grouping.

    Older versions generated many automatic keyword tags, which was confusing.
    From v27 onward, user-fed cases have one editable label. If the doctor does
    not provide one, it defaults to “用户新增”. Original Excel cases keep their
    source sheet as a non-editable group label.
    """
    label = _sanitize_case_label(getattr(rec, "sheet", "") if _is_user_case(rec) else getattr(rec, "sheet", ""))
    return [label or "用户新增"]


def _sanitize_case_label(label: str | None) -> str:
    """User-editable group label for doctor-fed cases.

    This is the label that appears in the knowledge-base group filter. It is
    separate from the fixed Excel source sheets. Empty labels fall back to
    “用户新增”.
    """
    text = str(label or "").strip()
    text = re.sub(r"[\r\n\t]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return "用户新增"
    return text[:24]


def _is_user_case(rec: CaseRecord) -> bool:
    return str(getattr(rec, "case_id", "")).startswith("USER-") or str(getattr(rec, "remarks", "")).find("由对话或上传文件加入") >= 0

def _normalize_legacy_user_labels() -> None:
    """Migrate old default display label to the new concise label.

    Only doctor-fed USER-* cases are changed; original Excel sheet names stay intact.
    """
    changed = False
    for rec in kb.records:
        if _is_user_case(rec) and str(getattr(rec, "sheet", "")).strip() in {"", "用户新增病例", "用户新增病例分组"}:
            rec.sheet = "用户新增"
            changed = True
    if changed:
        _save_user_cases()


def _apply_deleted_case_filter() -> None:
    deleted = _load_deleted_case_ids()
    if deleted:
        kb.records = [r for r in kb.records if r.case_id not in deleted]



def _knowledge_digest() -> Dict[str, Any]:
    """Small cached digest injected into normal AI answers.

    It avoids sending the whole database on every question while still giving
    the model a quick sense of what the local knowledge base contains.
    """
    try:
        source_mtime = DATA_PATH.stat().st_mtime if DATA_PATH.exists() else 0
        cached = json.loads(KB_DIGEST_PATH.read_text(encoding="utf-8")) if KB_DIGEST_PATH.exists() else {}
        if cached.get("source_mtime") == source_mtime and cached.get("case_count") == len(kb.records) and cached.get("article_count") == len(_load_articles()):
            return cached
    except Exception:
        pass
    sheets: Dict[str, int] = {}
    diagnosis_terms: Dict[str, int] = {}
    treatment_terms: Dict[str, int] = {}
    for r in kb.records:
        sheets[str(getattr(r, "sheet", "") or "未分组")] = sheets.get(str(getattr(r, "sheet", "") or "未分组"), 0) + 1
        for key in [getattr(r, "diagnosis", ""), getattr(r, "tnm", ""), getattr(r, "ls", "")]:
            k = str(key or "").strip()[:30]
            if k:
                diagnosis_terms[k] = diagnosis_terms.get(k, 0) + 1
        for key in [getattr(r, "surgery", ""), getattr(r, "other_treatment", "")]:
            k = str(key or "").strip()[:40]
            if k:
                treatment_terms[k] = treatment_terms.get(k, 0) + 1
    digest = {
        "updated_at": datetime.now().isoformat(timespec="seconds"),
        "source_mtime": DATA_PATH.stat().st_mtime if DATA_PATH.exists() else 0,
        "case_count": len(kb.records),
        "article_count": len(_load_articles()),
        "groups": sheets,
        "common_diagnosis_or_stage_terms": sorted(diagnosis_terms.items(), key=lambda x: x[1], reverse=True)[:18],
        "common_treatment_terms": sorted(treatment_terms.items(), key=lambda x: x[1], reverse=True)[:18],
    }
    try:
        _atomic_write_json(KB_DIGEST_PATH, digest)
    except Exception:
        pass
    return digest


def _load_library_snapshot() -> Dict[str, Any]:
    data = _read_json_with_backup(LIBRARY_STATE_PATH, {})
    return data if isinstance(data, dict) else {}


def _save_library_snapshot() -> None:
    snapshot = {
        "updated_at": datetime.now().isoformat(timespec="seconds"),
        "user_cases": _read_json_with_backup(USER_CASES_PATH, []),
        "articles": _read_json_with_backup(ARTICLES_PATH, []),
        "deleted_case_ids": sorted(_load_deleted_case_ids()),
    }
    _atomic_write_json(LIBRARY_STATE_PATH, snapshot)


def _restore_from_library_snapshot_if_needed() -> None:
    snap = _load_library_snapshot()
    if not snap:
        return
    if not _is_valid_json_file(USER_CASES_PATH):
        _atomic_write_json(USER_CASES_PATH, snap.get("user_cases") or [])
    if not _is_valid_json_file(ARTICLES_PATH):
        _atomic_write_json(ARTICLES_PATH, snap.get("articles") or [])
    if not _is_valid_json_file(DELETED_CASES_PATH):
        _atomic_write_json(DELETED_CASES_PATH, snap.get("deleted_case_ids") or [])


def _load_articles() -> List[Dict[str, Any]]:
    rows = _read_json_with_backup(ARTICLES_PATH, None)
    if isinstance(rows, list):
        return [_json_safe(x) for x in rows if isinstance(x, dict)]
    snap = _load_library_snapshot()
    rows = snap.get("articles") or []
    if isinstance(rows, list):
        return [_json_safe(x) for x in rows if isinstance(x, dict)]
    return []


def _atomic_write_json(path: Path, data: Any) -> None:
    """Safely persist JSON with backup and fsync.

    User-fed cases/articles must survive app restart and front-end updates.
    This writer keeps a .bak copy of the previous good file, writes to a
    temporary file, fsyncs it, then atomically replaces the target.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    payload = json.dumps(_json_safe(data), ensure_ascii=False, indent=2)
    backup = path.with_suffix(path.suffix + ".bak")
    if path.exists():
        try:
            shutil.copy2(path, backup)
        except Exception:
            pass
    tmp = path.with_suffix(path.suffix + f".{uuid.uuid4().hex}.tmp")
    with tmp.open("w", encoding="utf-8") as f:
        f.write(payload)
        f.flush()
        try:
            os.fsync(f.fileno())
        except Exception:
            pass
    tmp.replace(path)


def _read_json_with_backup(path: Path, default: Any = None) -> Any:
    """Read the main JSON file first and use its backup only if needed.

    An empty list or dict is valid persisted state. Older versions treated it
    like a broken file and loaded stale backup rows, which could resurrect
    deleted cases and then hide them again through deletion tombstones.
    """
    for cand in [path, path.with_suffix(path.suffix + ".bak")]:
        try:
            if cand.exists():
                return json.loads(cand.read_text(encoding="utf-8"))
        except Exception:
            continue
    return default


def _is_valid_json_file(path: Path) -> bool:
    try:
        json.loads(path.read_text(encoding="utf-8"))
        return True
    except Exception:
        return False


def _read_nonempty_json(path: Path) -> Any | None:
    data = _read_json_with_backup(path, None)
    return data if data not in (None, [], {}) else None


def _legacy_json_candidates(filename: str) -> List[Path]:
    candidates = [PROJECT_DATA_DIR / filename]
    # When a new version is unzipped next to an older project folder, import the
    # old data automatically on first run. This is intentionally shallow to avoid
    # scanning the whole computer.
    try:
        for p in BASE_DIR.parent.glob("*/data/" + filename):
            if p not in candidates:
                candidates.append(p)
    except Exception:
        pass
    return candidates


def _migrate_project_json_if_needed(target: Path, legacy: Path, default: Any) -> None:
    """Initialize persistent storage once, without resurrecting deleted data.

    The persistent directory is the source of truth. Packaged data/*.json files
    are only used on the first run when no persistent file exists. If articles
    have ever been deliberately deleted to empty, a marker prevents later
    package updates from re-importing old articles.
    """
    target.parent.mkdir(parents=True, exist_ok=True)
    if target.exists():
        return
    if target.name == "articles.json" and ARTICLES_DELETED_MARKER.exists():
        _atomic_write_json(target, [])
        return
    filename = legacy.name
    best = None
    best_len = -1
    for cand in _legacy_json_candidates(filename):
        # Ignore packaged empty defaults and prefer the largest previous user file.
        data = _read_nonempty_json(cand)
        if data is not None:
            score = len(data) if hasattr(data, "__len__") else 1
            if score > best_len:
                best, best_len = data, score
    if best is not None:
        _atomic_write_json(target, best)
        return
    _atomic_write_json(target, default)



def _merge_list_json_by_key(target: Path, filename: str, key_names: List[str]) -> None:
    """Merge legacy JSON rows into the persistent file even when the persistent
    file already exists. This prevents data loss across front-end updates.
    """
    rows: List[Dict[str, Any]] = []
    seen: set[str] = set()

    def row_key(row: Dict[str, Any]) -> str:
        for k in key_names:
            val = str(row.get(k) or "").strip()
            if val:
                return f"{k}:{val}"
        # Fallback content signature.
        compact = json.dumps(_json_safe(row), ensure_ascii=False, sort_keys=True)
        return "hash:" + hashlib.sha1(compact.encode("utf-8", errors="ignore")).hexdigest()

    for source in [target] + _legacy_json_candidates(filename):
        data = _read_nonempty_json(source)
        if not isinstance(data, list):
            continue
        for item in data:
            if not isinstance(item, dict):
                continue
            key = row_key(item)
            if key in seen:
                continue
            seen.add(key)
            rows.append(item)
    if rows:
        _atomic_write_json(target, rows)


def _merge_dict_json(target: Path, filename: str) -> None:
    merged: Dict[str, Any] = {}
    for source in _legacy_json_candidates(filename) + [target]:
        data = _read_nonempty_json(source)
        if isinstance(data, dict):
            merged.update(data)
    if merged:
        _atomic_write_json(target, merged)


def _merge_deleted_ids(target: Path, filename: str) -> None:
    ids: set[str] = set()
    for source in [target] + _legacy_json_candidates(filename):
        data = _read_nonempty_json(source)
        if isinstance(data, list):
            ids.update(str(x) for x in data if str(x).strip())
    if ids:
        _atomic_write_json(target, sorted(ids))

def _migrate_uploads_if_needed() -> None:
    legacy_uploads = BASE_DIR / "uploads"
    if not legacy_uploads.exists() or legacy_uploads.resolve() == UPLOAD_DIR.resolve():
        return
    for src in legacy_uploads.iterdir():
        if src.is_file() and src.name != ".gitkeep":
            dst = UPLOAD_DIR / src.name
            if not dst.exists():
                try:
                    shutil.copy2(src, dst)
                except Exception:
                    pass


def _file_sha256(path: Path) -> str:
    try:
        if not path.exists():
            return ""
        h = hashlib.sha256()
        with path.open("rb") as f:
            for chunk in iter(lambda: f.read(1024 * 1024), b""):
                h.update(chunk)
        return h.hexdigest()
    except Exception:
        return ""


def _ensure_json_file(path: Path, default: Any) -> None:
    """Create an empty persistent JSON file if absent.

    v30 change: packaged project JSON files are no longer treated as runtime
    storage. This fixes the recurring “0 个病例、70 篇文章” problem caused by
    re-importing bundled demo JSON after every front-end replacement.
    """
    if _is_valid_json_file(path):
        return
    _atomic_write_json(path, _read_json_with_backup(path, default))


def _reconcile_user_case_tombstones() -> int:
    """Restore saved user cases that old versions accidentally hid on restart.

    A USER-* row present in user_cases.json is active persisted data. Deleting a
    user case removes that row, so a simultaneous tombstone is a stale conflict.
    Built-in Excel case tombstones remain untouched.
    """
    rows = _read_json_with_backup(USER_CASES_PATH, [])
    if not isinstance(rows, list):
        return 0
    saved_ids = {
        str(row.get("case_id") or "").strip()
        for row in rows
        if isinstance(row, dict) and str(row.get("case_id") or "").startswith("USER-")
    }
    deleted = _load_deleted_case_ids()
    restored = sorted(saved_ids & deleted)
    if not restored:
        return 0
    _atomic_write_json(DELETED_CASES_PATH, sorted(deleted - set(restored)))
    _atomic_write_json(PERSISTENCE_REPAIR_LOG_PATH, {
        "updated_at": datetime.now().isoformat(timespec="seconds"),
        "repair": "restored_saved_user_cases_hidden_by_tombstones",
        "restored_user_case_count": len(restored),
        "restored_user_case_ids": restored,
    })
    return len(restored)


def _deduplicate_saved_user_cases() -> int:
    """Keep the oldest saved copy of each clinically identical user case."""
    rows = _read_json_with_backup(USER_CASES_PATH, [])
    if not isinstance(rows, list):
        return 0

    def sort_key(row: Dict[str, Any]) -> tuple[int, str]:
        case_id = str(row.get("case_id") or "")
        try:
            return int(case_id.split("-", 1)[1]), case_id
        except Exception:
            return 10**12, case_id

    kept: List[Dict[str, Any]] = []
    removed_ids: List[str] = []
    seen: set[str] = set()
    signatures_updated = False
    for row in sorted((x for x in rows if isinstance(x, dict)), key=sort_key):
        fingerprint = _case_duplicate_fingerprint(row)
        if fingerprint and fingerprint in seen:
            case_id = str(row.get("case_id") or "").strip()
            if case_id:
                removed_ids.append(case_id)
            continue
        if fingerprint:
            seen.add(fingerprint)
            if str(row.get("case_signature") or "") != fingerprint:
                row = {**row, "case_signature": fingerprint}
                signatures_updated = True
        kept.append(row)
    if not removed_ids and not signatures_updated:
        return 0
    _atomic_write_json(USER_CASES_PATH, kept)
    if not removed_ids:
        return 0
    deleted = _load_deleted_case_ids()
    deleted.update(removed_ids)
    _atomic_write_json(DELETED_CASES_PATH, sorted(deleted))
    _atomic_write_json(PERSISTENCE_REPAIR_LOG_PATH, {
        "updated_at": datetime.now().isoformat(timespec="seconds"),
        "repair": "removed_duplicate_user_cases",
        "kept_user_case_count": len(kept),
        "removed_duplicate_count": len(removed_ids),
        "removed_duplicate_case_ids": removed_ids,
    })
    return len(removed_ids)


def _clear_packaged_articles_if_resurrected() -> None:
    """If old packaged demo articles were copied into persistent storage, clear them.

    This only triggers when persistent articles.json is byte-for-byte equivalent
    to the bundled data/articles.json, i.e. it is almost certainly the old fixed
    package content rather than the user's own article library.
    """
    packaged = PROJECT_DATA_DIR / "articles.json"
    if not ARTICLES_PATH.exists() or not packaged.exists():
        return
    try:
        rows = json.loads(ARTICLES_PATH.read_text(encoding="utf-8"))
        if not isinstance(rows, list) or len(rows) == 0:
            return
        if _file_sha256(ARTICLES_PATH) and _file_sha256(ARTICLES_PATH) == _file_sha256(packaged):
            _atomic_write_json(ARTICLES_PATH, [])
            ARTICLES_DELETED_MARKER.write_text(datetime.now().isoformat(timespec="seconds"), encoding="utf-8")
    except Exception:
        pass


def _clear_v34_requested_article_library() -> None:
    """User requested clearing the 70-article library. Run once for v34.

    This clears persistent articles and writes the delete marker so bundled
    demo/default articles will not resurrect on restart.
    """
    if ARTICLES_V34_CLEARED_MARKER.exists():
        return
    rows = _load_articles()
    if rows:
        _atomic_write_json(ARTICLES_PATH, [])
    try:
        ARTICLES_DELETED_MARKER.write_text(datetime.now().isoformat(timespec="seconds"), encoding="utf-8")
        ARTICLES_V34_CLEARED_MARKER.write_text(datetime.now().isoformat(timespec="seconds"), encoding="utf-8")
    except Exception:
        pass


def _ensure_storage_files() -> None:
    """Prepare persistent storage.

    Runtime user data lives only in ~/.uscc_scc_flask_data unless USCC_DATA_DIR
    is explicitly set. Bundled JSON files are never re-imported automatically;
    this prevents deleted articles from reappearing after an update.
    """
    _restore_from_library_snapshot_if_needed()
    _ensure_json_file(USER_CASES_PATH, [])
    _ensure_json_file(DELETED_CASES_PATH, [])
    _ensure_json_file(ARTICLES_PATH, [])
    _ensure_json_file(CASE_TAGS_PATH, {})
    _clear_packaged_articles_if_resurrected()
    _reconcile_user_case_tombstones()
    _deduplicate_saved_user_cases()
    _migrate_uploads_if_needed()
    _save_library_snapshot()
    state = {
        "storage_version": "v36",
        "updated_at": datetime.now().isoformat(timespec="seconds"),
        "data_dir": str(PERSISTENT_DATA_DIR),
        "source_of_truth": "persistent_only",
        "note": "Bundled data/*.json is ignored at runtime; user cases/articles are loaded only from the persistent data directory.",
    }
    _atomic_write_json(MIGRATION_STATE_PATH, state)
    _write_storage_manifest()


def _save_articles(rows: List[Dict[str, Any]]) -> None:
    rows = [_json_safe(x) for x in (rows or []) if isinstance(x, dict)]
    _atomic_write_json(ARTICLES_PATH, rows)
    if rows:
        try:
            ARTICLES_DELETED_MARKER.unlink(missing_ok=True)
        except Exception:
            pass
    else:
        try:
            ARTICLES_DELETED_MARKER.write_text(datetime.now().isoformat(timespec="seconds"), encoding="utf-8")
        except Exception:
            pass
    _save_library_snapshot()
    _write_storage_manifest()


def _write_storage_manifest() -> None:
    try:
        case_count = len(json.loads(USER_CASES_PATH.read_text(encoding="utf-8"))) if USER_CASES_PATH.exists() else 0
    except Exception:
        case_count = -1
    try:
        article_count = len(json.loads(ARTICLES_PATH.read_text(encoding="utf-8"))) if ARTICLES_PATH.exists() else 0
    except Exception:
        article_count = -1
    _atomic_write_json(PERSISTENT_DATA_DIR / "storage_manifest.json", {
        "updated_at": datetime.now().isoformat(timespec="seconds"),
        "storage_version": "v36",
        "data_dir": str(PERSISTENT_DATA_DIR),
        "user_cases_path": str(USER_CASES_PATH),
        "user_case_count": case_count,
        "articles_path": str(ARTICLES_PATH),
        "article_count": article_count,
        "articles_deleted_marker": ARTICLES_DELETED_MARKER.exists(),
        "deleted_cases_path": str(DELETED_CASES_PATH),
        "uploads_path": str(UPLOAD_DIR),
        "persistence_repair_log_path": str(PERSISTENCE_REPAIR_LOG_PATH),
    })


def _next_article_id() -> str:
    nums = []
    for a in _load_articles():
        aid = str(a.get("article_id") or "")
        if aid.startswith("ARTICLE-"):
            try:
                nums.append(int(aid.split("-", 1)[1]))
            except Exception:
                pass
    return f"ARTICLE-{(max(nums) if nums else 0) + 1:03d}"


def _article_matches_query(article: Dict[str, Any], query: str) -> bool:
    q = (query or "").strip().lower()
    if not q:
        return True
    haystack = "\n".join(str(article.get(k) or "") for k in ["title", "authors", "journal", "year", "doi", "source_url", "keywords", "abstract", "content", "notes"]).lower()
    terms = [x for x in q.replace("，", " ").replace(",", " ").split() if x]
    return all(t in haystack for t in terms) if terms else q in haystack


def _article_terms(query: str) -> List[str]:
    q = (query or "").lower()
    raw_terms = [x for x in re.split(r"[\s,，。；;:：/\\()（）\[\]{}<>]+", q) if len(x.strip()) >= 2]
    domain_terms = []
    keyword_map = {
        "scc": ["scc", "鳞状细胞癌", "鳞癌", "squamous"],
        "ls": ["ls", "硬化性苔藓", "lichen sclerosus"],
        "尿道": ["尿道", "urethral", "urethra"],
        "狭窄": ["狭窄", "stricture"],
        "手术": ["手术", "切除", "尿道切除", "阴茎切除", "会阴尿道造口", "surgery", "excision", "urethrectomy", "penectomy"],
        "治疗": ["治疗", "用药", "化疗", "放疗", "免疫治疗", "therapy", "treatment", "chemotherapy", "radiotherapy", "immunotherapy"],
        "复发": ["复发", "转移", "淋巴结", "recurrence", "metastasis", "lymph node"],
        "病理": ["病理", "活检", "免疫组化", "p63", "p40", "ck5", "ki-67", "pathology", "biopsy"],
    }
    for trigger, terms in keyword_map.items():
        if any(t in q for t in terms):
            domain_terms.extend(terms)
    # Always keep core disease anchors so short Chinese病例描述也能命中文献。
    if any(t in q for t in ["尿道", "scc", "鳞癌", "鳞状细胞癌"]):
        domain_terms.extend(["尿道", "urethral", "scc", "鳞状细胞癌", "鳞癌"])
    seen, out = set(), []
    for t in raw_terms + domain_terms:
        t = t.strip().lower()
        if t and t not in seen:
            seen.add(t); out.append(t)
    return out[:60]


def _extract_article_value_points(article: Dict[str, Any], terms: List[str], max_points: int = 2) -> List[str]:
    """Return concise Chinese clinical takeaways, not long copied excerpts.

    This is deliberately rule-based: it turns matched literature into short
    reminders for doctors and keeps the full text available through article_url.
    """
    text = "\n".join(str(article.get(k) or "") for k in ["title", "keywords", "abstract", "content", "notes"]).lower()
    points: List[str] = []

    def add(point: str) -> None:
        if point and point not in points and len(points) < max_points:
            points.append(point)

    if any(x in text for x in ["lichen sclerosus", "硬化性苔藓", " ls ", "ls,"]):
        add("若当前病例存在 LS 或长期尿道狭窄，可重点参考其癌前病变识别、活检时机与长期随访策略。")
    if any(x in text for x in ["surgery", "urethrectomy", "penectomy", "excision", "手术", "切除", "尿道切除", "阴茎切除", "会阴尿道造口"]):
        add("文献对局部切除、尿道切除或更扩大手术边界有参考价值，需结合肿瘤部位、浸润深度和功能保留目标讨论。")
    if any(x in text for x in ["chemotherapy", "radiotherapy", "immunotherapy", "cisplatin", "放疗", "化疗", "免疫治疗", "顺铂"]):
        add("涉及放化疗或系统治疗，可用于讨论局部晚期、复发转移或围手术期综合治疗的备选路径。")
    if any(x in text for x in ["lymph", "node", "metastasis", "淋巴结", "转移"]):
        add("涉及淋巴结或转移风险，可提示完善影像分期、腹股沟/盆腔淋巴结评估与随访重点。")
    if any(x in text for x in ["recurrence", "follow-up", "survival", "复发", "随访", "预后"]):
        add("涉及复发或预后，可参考其随访间隔、复查项目和复发后处理思路。")
    if any(x in text for x in ["pathology", "biopsy", "p40", "p63", "ck5", "ki-67", "病理", "活检", "免疫组化"]):
        add("涉及病理或免疫组化，可用于核对活检充分性、SCC 证据链和鉴别诊断要点。")

    if not points:
        hit = "、".join([t for t in terms[:4] if len(t) >= 2])
        add(f"与当前问题存在关键词关联{('：' + hit) if hit else ''}，建议点开详情查看原文适用范围。")
    return points[:max_points]


def _search_articles(query: str, limit: int = 5) -> List[Dict[str, Any]]:
    articles = _load_articles()
    if not articles:
        return []
    terms = _article_terms(query)
    scored = []
    for a in articles:
        text = "\n".join(str(a.get(k) or "") for k in ["title", "keywords", "abstract", "content", "notes"]).lower()
        score = 0
        hit_terms = []
        for term in terms:
            c = text.count(term)
            if c:
                hit_terms.append(term)
                score += c * (3 if term in {"尿道", "urethral", "scc", "鳞癌", "鳞状细胞癌", "治疗", "手术"} else 1)
        if score > 0:
            item = dict(a)
            item["match_score"] = score
            item["hit_terms"] = hit_terms[:12]
            if len(str(item.get("content") or "")) > 700:
                item["content_preview"] = str(item.get("content"))[:700] + "..."
            else:
                item["content_preview"] = str(item.get("content") or "")
            item["article_url"] = f"/article/{item.get('article_id')}"
            item["value_points"] = _extract_article_value_points(item, terms)
            scored.append(item)
    scored.sort(key=lambda x: (x.get("match_score", 0), x.get("created_at", "")), reverse=True)
    if not scored:
        return []
    best = int(scored[0].get("match_score", 0) or 0)
    # v19: 至少给出 2 篇可核对文献，至多 4 篇；但不机械取满 4 篇。
    # 若最佳匹配明显高于其他文章，只展示 2 篇；若多篇接近，再展示到最多 4 篇。
    threshold = max(2, int(best * 0.55))
    filtered = [x for x in scored if int(x.get("match_score", 0) or 0) >= threshold]
    if len(filtered) < min(2, len(scored)):
        filtered = scored[:min(2, len(scored))]
    return _json_safe(filtered[:limit])


def _extract_article_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if suffix == ".pdf":
        # Prefer pypdf; fall back to PyPDF2 if the environment already has it.
        try:
            from pypdf import PdfReader
        except Exception:
            try:
                from PyPDF2 import PdfReader
            except Exception as exc:
                raise RuntimeError("PDF 解析需要安装 pypdf：pip install pypdf") from exc
        reader = PdfReader(str(path))
        parts = []
        for i, page in enumerate(reader.pages[:80], 1):
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            if txt.strip():
                parts.append(f"# Page {i}\n{txt.strip()}")
        return "\n\n".join(parts)
    if suffix in {".xlsx", ".xls"}:
        import pandas as pd
        xl = pd.ExcelFile(path)
        parts = []
        for sheet in xl.sheet_names:
            df = pd.read_excel(path, sheet_name=sheet, dtype=str).fillna("")
            parts.append(f"# {sheet}")
            parts.append(df.to_csv(index=False, sep="\t"))
        return "\n".join(parts)
    if suffix == ".csv":
        import pandas as pd
        df = pd.read_csv(path, dtype=str).fillna("")
        return df.to_csv(index=False, sep="\t")
    if suffix == ".md":
        return path.read_text(encoding="utf-8", errors="ignore")
    return ""




def _extract_article_images(path: Path, max_images: int = 8) -> List[Dict[str, Any]]:
    """Extract embedded article images to persistent uploads for optional multimodal analysis.

    DOCX images are read from word/media. PDF images use PyMuPDF when installed;
    if PyMuPDF is absent, text parsing still succeeds and the note tells the user.
    """
    suffix = path.suffix.lower()
    images: List[Dict[str, Any]] = []

    def save_bytes(data: bytes, suffix_hint: str, label: str) -> None:
        if len(images) >= max_images or not data:
            return
        ext = suffix_hint.lower()
        if ext not in {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp", ".tif", ".tiff"}:
            ext = ".png"
        name = f"{uuid.uuid4().hex}_article_image{ext}"
        dest = UPLOAD_DIR / name
        dest.write_bytes(data)
        images.append({"type": "image", "filename": label or name, "stored_as": name, "url": f"/uploads/{name}"})

    try:
        if suffix == ".docx":
            with zipfile.ZipFile(path) as zf:
                for info in zf.infolist():
                    if len(images) >= max_images:
                        break
                    if not info.filename.lower().startswith("word/media/"):
                        continue
                    ext = Path(info.filename).suffix or ".png"
                    save_bytes(zf.read(info), ext, Path(info.filename).name)
        elif suffix == ".pdf":
            try:
                import fitz  # PyMuPDF
            except Exception:
                return images
            doc = fitz.open(str(path))
            for page_index in range(min(len(doc), 20)):
                if len(images) >= max_images:
                    break
                for img_index, img in enumerate(doc.get_page_images(page_index, full=True)):
                    if len(images) >= max_images:
                        break
                    xref = img[0]
                    base = doc.extract_image(xref)
                    ext = "." + str(base.get("ext") or "png")
                    save_bytes(base.get("image") or b"", ext, f"page{page_index+1}_image{img_index+1}{ext}")
    except Exception:
        return images
    return images


def _summarize_article_images_for_doctor(images: List[Dict[str, Any]], filename: str, article_text: str = "", api_key: str = "", provider: str = "", model: str = "", base_url: str = "") -> str:
    if not images:
        return ""
    if not api_key:
        return f"\n\n【文章图片】检测到 {len(images)} 张图片。未配置 API Key，已保存图片，但未进行图像内容分析。"
    context = (article_text or "")[:6000]
    prompt = (
        f"请结合文章文件《{filename}》的正文和其中提取出的图片进行医学阅读。\n"
        "任务：1）逐图用中文概括可观察内容；2）结合正文判断图片支持了文章的哪些临床经验；"
        "3）提炼对男性尿道SCC、LS相关恶变、病理/影像/手术或随访有价值的经验；"
        "4）说明哪些内容需要医生打开原图或原文再次核对。不要编造无法从图中或正文中看出的诊断。\n\n"
        f"文章正文节选：\n{context}"
    )
    result = ask_llm(
        question=prompt,
        report={"similar_cases": [], "related_articles": [], "treatment_outcomes": {}, "risk": {"missing_items": []}, "knowledge_digest": _knowledge_digest()},
        patient={},
        history=[],
        attachments=images,
        api_key_override=api_key,
        model_override=model,
        provider_override=provider or "openai",
        base_url_override=base_url,
        mode_override="article_image_analysis",
    )
    answer = str(result.get("answer") or "").strip()
    if result.get("provider", "").startswith("local_fallback"):
        return f"\n\n【文章图片】检测到 {len(images)} 张图片，但图像分析未成功：{result.get('error') or '未调用到可用多模态模型'}"
    return f"\n\n【文章图片与正文联合分析】\n{answer}" if answer else f"\n\n【文章图片】检测到 {len(images)} 张图片，但模型未返回有效图像分析。"


def _article_signature_fields(fields: Dict[str, Any]) -> str:
    title = str(fields.get("title") or "").strip().lower()
    doi = _normalize_article_doi(fields.get("doi") or "")
    source_url = _normalize_article_url(fields.get("source_url") or "")
    source_file = str(fields.get("source_file") or "").strip().lower()
    content = str(fields.get("content") or fields.get("abstract") or "").strip().lower()
    normalized = re.sub(r"\s+", "", title + "|" + doi + "|" + source_url + "|" + source_file + "|" + content[:80000])
    return hashlib.sha256(normalized.encode("utf-8", errors="ignore")).hexdigest() if normalized else ""


def _normalize_article_doi(value: Any) -> str:
    text = str(value or "").strip().lower()
    text = re.sub(r"^https?://(?:dx\.)?doi\.org/", "", text)
    match = re.search(r"10\.\d{4,9}/[-._;()/:a-z0-9]+", text, re.I)
    return match.group(0).rstrip(".,;") if match else ""


def _normalize_article_url(value: Any) -> str:
    text = str(value or "").strip().lower()
    return re.sub(r"[?#].*$", "", text).rstrip("/")


def _normalize_article_title(value: Any) -> str:
    return re.sub(r"[\W_]+", "", str(value or "").strip().lower(), flags=re.UNICODE)


def _find_duplicate_article(fields: Dict[str, Any], exclude_id: str = "") -> Dict[str, Any] | None:
    sig = _article_signature_fields(fields)
    doi = _normalize_article_doi(fields.get("doi") or fields.get("source_url") or "")
    source_url = _normalize_article_url(fields.get("source_url") or "")
    title = _normalize_article_title(fields.get("title") or "")
    source_file = re.sub(r"\s+", "", str(fields.get("source_file") or "").lower())
    content_head = re.sub(r"\s+", "", str(fields.get("content") or fields.get("abstract") or "").lower())[:2000]
    for row in _load_articles():
        if exclude_id and str(row.get("article_id")) == exclude_id:
            continue
        row_sig = row.get("signature") or _article_signature_fields(row)
        if sig and row_sig == sig:
            return row
        row_doi = _normalize_article_doi(row.get("doi") or row.get("source_url") or "")
        row_url = _normalize_article_url(row.get("source_url") or "")
        row_title = _normalize_article_title(row.get("title") or "")
        row_source = re.sub(r"\s+", "", str(row.get("source_file") or "").lower())
        row_head = re.sub(r"\s+", "", str(row.get("content") or row.get("abstract") or "").lower())[:2000]
        if doi and row_doi and doi == row_doi:
            return row
        if source_url and row_url and source_url == row_url:
            return row
        if len(title) >= 12 and title == row_title:
            return row
        if source_file and row_source and source_file == row_source:
            return row
        if title and row_title and title == row_title and content_head and row_head and content_head == row_head:
            return row
    return None

def _add_article(fields: Dict[str, Any]) -> Dict[str, Any]:
    title = str(fields.get("title") or "").strip() or "未命名文献"
    content = str(fields.get("content") or "").strip()
    abstract = str(fields.get("abstract") or "").strip()
    if not content and not abstract:
        raise ValueError("请提供文章正文或摘要。")
    base = {
        "title": title,
        "authors": str(fields.get("authors") or "").strip(),
        "journal": str(fields.get("journal") or "").strip(),
        "year": str(fields.get("year") or "").strip(),
        "doi": _normalize_article_doi(fields.get("doi") or fields.get("source_url") or ""),
        "keywords": str(fields.get("keywords") or "").strip(),
        "abstract": abstract,
        "content": content,
        "notes": str(fields.get("notes") or "").strip(),
        "source_file": str(fields.get("source_file") or "").strip(),
        "source_url": str(fields.get("source_url") or "").strip(),
        "stored_as": str(fields.get("stored_as") or "").strip(),
        "article_images": fields.get("article_images") or [],
        "created_at": str(fields.get("created_at") or datetime.now().isoformat(timespec="seconds")),
    }
    duplicate = _find_duplicate_article(base)
    if duplicate:
        dup = dict(duplicate)
        dup["duplicate"] = True
        dup["message"] = f"检测到相同或高度重复文章：{duplicate.get('article_id')}，已跳过重复收录。"
        return _json_safe(dup)
    article = {"article_id": _next_article_id(), **base, "signature": _article_signature_fields(base)}
    rows = _load_articles()
    rows.append(article)
    _save_articles(rows)
    return article


def _extract_row_field(row: Dict[str, Any], aliases: List[str]) -> str:
    norm = lambda s: re.sub(r"[\s_\-:：()（）]+", "", str(s or "").lower())
    row_norm = {norm(k): v for k, v in row.items()}
    for a in aliases:
        ak = norm(a)
        for k, v in row_norm.items():
            if ak and (ak == k or ak in k):
                val = _json_safe(v)
                if str(val).strip() and str(val).lower() not in {"nan", "none", "null"}:
                    return str(val).strip()
    return ""


def _row_to_candidate(row: Dict[str, Any], idx: int, source_name: str, batch_id: str) -> Dict[str, Any]:
    aliases = {
        "name": ["姓名", "患者", "患者姓名", "name"],
        "sex": ["性别", "sex", "gender"],
        "age": ["年龄", "年龄（岁）", "年龄(岁)", "age"],
        "diagnosis": ["诊断", "主要诊断", "疾病名称/诊断", "病理诊断"],
        "ls": ["LS", "硬化性苔藓", "lichen sclerosus"],
        "tnm": ["TNM", "分期", "TNM分期"],
        "lymph_node": ["淋巴结", "淋巴转移", "有无淋巴结转移"],
        "history": ["病史", "病史要点", "现病史", "既往史", "主诉"],
        "symptoms": ["症状", "体征", "临床表现", "症状/体征"],
        "tumor": ["肿瘤", "肿物", "肿块", "占位", "肿瘤情况", "病灶位置"],
        "imaging": ["影像", "影像学", "CT", "MRI", "尿道镜", "检查所见"],
        "pathology": ["病理", "活检", "免疫组化", "病理信息", "病理诊断"],
        "surgery": ["手术", "术式", "手术治疗", "治疗方式"],
        "other_treatment": ["其他治疗", "用药", "化疗", "放疗", "免疫治疗"],
        "followup": ["随访", "转归", "预后", "结局", "复发"],
    }
    item = {k: _extract_row_field(row, v) for k, v in aliases.items()}
    blob = "；".join(f"{k}:{v}" for k, v in row.items() if str(v).strip() and str(v).lower() != "nan")
    item.update({
        "candidate_id": f"{batch_id}-{idx+1:04d}",
        "batch_id": batch_id,
        "row_index": idx,
        "source_file": source_name,
        "free_text": blob[:5000],
        "display_title": "｜".join([x for x in [item.get("name"), (item.get("age") + "岁") if item.get("age") else "", item.get("sex"), item.get("diagnosis") or item.get("pathology")] if x])[:120] or f"候选病例 {idx+1}",
    })
    return _json_safe(item)



def _looks_like_header_or_prompt_row(row: Dict[str, Any]) -> bool:
    """Detect rows that are prompts/headers rather than real patient records.

    The common Excel structure is one row of group headers and one row of
    column prompts such as “姓名/性别/年龄/病理诊断”. Earlier versions could treat
    that prompt row as a patient. This function is intentionally conservative:
    if most cells are column names or empty labels and there is no concrete
    patient value, the row is skipped.
    """
    values = [str(v).strip() for v in row.values() if str(v).strip() and str(v).strip().lower() not in {"nan", "none", "null", "unnamed"}]
    if not values:
        return True
    joined = " ".join(values)
    normalized_values = [re.sub(r"[\s:：()（）/\\_-]+", "", v).lower() for v in values]
    header_tokens = [
        "姓名", "患者", "患者姓名", "性别", "年龄", "年龄岁", "入院日期", "疾病名称诊断", "诊断",
        "病史要点", "临床症状信息", "症状体征", "肿瘤情况", "肿瘤分化程度grade分级", "tnm分期grade分级",
        "有无淋巴结转移", "有无硬化性苔藓ls", "免疫", "免疫组化", "影像信息", "病理诊断",
        "治疗方案", "手术治疗", "手术时间", "其他治疗", "术后影像", "术后复发转移情况", "复发转移后处理",
        "随访结果", "备注", "基本信息", "住院号", "病案号", "既往尿道重建手术",
        "疾病名称", "临床症状", "症状信息", "grade分级"
    ]
    header_norm = [re.sub(r"[\s:：()（）/\\_-]+", "", x).lower() for x in header_tokens]
    label_like = 0
    for v in normalized_values:
        if not v:
            continue
        if v.startswith("unnamed") or any(v == h or v in h or h in v for h in header_norm):
            label_like += 1
    label_ratio = label_like / max(len(normalized_values), 1)

    # Concrete values that make a row look like an actual patient, not a prompt row.
    concrete_patterns = [
        r"\d{1,3}\s*岁", r"^\d{1,3}$", r"\d{4}[-/]\d{1,2}[-/]\d{1,2}",
        r"T\s*[0-4][a-c]?\s*N\s*[0-3xX]?\s*M\s*[01xX]?", r"G[1-4]",
        r"鳞癌|鳞状细胞癌|SCC|尿道狭窄|排尿困难|血尿|转移|复发|紫杉醇|顺铂|放疗|切除术|造瘘|活检|浸润",
    ]
    concrete_count = 0
    for v in values:
        vv = str(v).strip()
        # 纯字段名、合并表头、斜杠占位符不算真实病例内容。
        if vv in {"/", "\\", "-", "—"}:
            continue
        if vv in {"男", "女", "男性", "女性"}:
            concrete_count += 1
        elif any(re.search(pat, vv, re.I) for pat in concrete_patterns):
            concrete_count += 1

    # 如果整行几乎全是字段提示词，即使存在“疾病名称/诊断”这类词，也判定为表头。
    if label_ratio >= 0.30 and label_like >= 4 and concrete_count <= 1:
        return True
    if label_ratio >= 0.45 and concrete_count == 0:
        return True
    if label_like >= 5 and concrete_count <= 1:
        return True
    # A row composed of short labels is not a case even if it contains words like “病理/手术/随访”.
    if len(values) >= 3 and label_ratio >= 0.35 and sum(len(v) <= 12 for v in values) >= len(values) * 0.65 and concrete_count == 0:
        return True
    return False


def _values_look_like_header(vals: List[Any]) -> bool:
    row = {f"c{i}": v for i, v in enumerate(vals)}
    nonempty = [str(v).strip() for v in vals if str(v).strip() and str(v).strip().lower() not in {"nan", "none", "null"}]
    if len(nonempty) < 2:
        return False
    joined = " ".join(nonempty)
    words = ["姓名", "患者", "性别", "年龄", "诊断", "病理", "影像", "治疗", "手术", "随访", "住院号", "病案号", "症状", "肿瘤", "淋巴结"]
    return sum(1 for w in words if w in joined) >= 2 or _looks_like_header_or_prompt_row(row)


def _merge_header_labels(top: List[Any], second: List[Any]) -> List[str]:
    headers: List[str] = []
    for i, (a, b) in enumerate(zip(top, second)):
        aa = str(a or "").strip()
        bb = str(b or "").strip()
        # Pandas creates "Unnamed: x" for merged cells; those should not become labels.
        if aa.lower().startswith("unnamed"):
            aa = ""
        if bb.lower().startswith("unnamed"):
            bb = ""
        if bb:
            headers.append(bb)
        elif aa:
            headers.append(aa)
        else:
            headers.append(f"字段{i+1}")
    return headers


def _append_candidate_from_row(row_dict: Dict[str, Any], source_name: str, batch_id: str, candidates: List[Dict[str, Any]], max_rows: int) -> None:
    vals = [str(v).strip() for v in row_dict.values() if str(v).strip() and str(v).strip().lower() not in {"nan", "none", "null"}]
    if len(vals) < 2:
        return
    if _looks_like_header_or_prompt_row(row_dict):
        return
    candidates.append(_row_to_candidate(row_dict, len(candidates), source_name, batch_id))


def _dataframe_candidates_from_sheet(df, source_name: str, batch_id: str, candidates: List[Dict[str, Any]], max_rows: int) -> None:
    import pandas as pd
    df = df.fillna("")
    if df.empty or len(candidates) >= max_rows:
        return

    header_cols = [str(c).strip() for c in df.columns]
    known_cols = sum(1 for c in header_cols if any(w in c for w in ["姓名", "患者", "性别", "年龄", "诊断", "病理", "影像", "治疗", "手术", "随访", "住院号", "病案号"]))

    # Case A: pandas used the first group-header row as column names, and the first
    # data row is the actual prompt row. Promote that first data row to headers and
    # parse rows after it. This fixes two-line headers like the uploaded USCC sheet.
    if len(df) >= 2:
        first_values = [str(x).strip() for x in df.iloc[0].tolist()]
        if _values_look_like_header(first_values):
            headers = _merge_header_labels(header_cols, first_values)
            body = df.iloc[1:].copy()
            body.columns = headers[:len(body.columns)]
            for _, row in body.head(max_rows - len(candidates)).iterrows():
                _append_candidate_from_row({str(k): v for k, v in row.to_dict().items()}, source_name, batch_id, candidates, max_rows)
                if len(candidates) >= max_rows:
                    return
            return

    # Case B: ordinary table: the Excel first row already became useful columns.
    if known_cols >= 2:
        for _, row in df.head(max_rows - len(candidates)).iterrows():
            _append_candidate_from_row({str(k): v for k, v in row.to_dict().items()}, source_name, batch_id, candidates, max_rows)
            if len(candidates) >= max_rows:
                return
        return

    # Case C: headerless sheet. Locate a prompt row, optionally merging with the row
    # above if it contains group labels, then parse only subsequent rows.
    raw = df.copy()
    for ridx in range(min(len(raw), 12)):
        vals = [str(x).strip() for x in raw.iloc[ridx].tolist()]
        if not _values_look_like_header(vals):
            continue
        top = [str(x).strip() for x in raw.iloc[ridx-1].tolist()] if ridx > 0 else [""] * len(vals)
        headers = _merge_header_labels(top, vals)
        body = raw.iloc[ridx+1:].copy()
        body.columns = headers[:len(body.columns)]
        for _, row in body.head(max_rows - len(candidates)).iterrows():
            _append_candidate_from_row({str(k): v for k, v in row.to_dict().items()}, source_name, batch_id, candidates, max_rows)
            if len(candidates) >= max_rows:
                return
        return

def _extract_case_batch_candidates(path: Path, original_name: str, max_rows: int = 1200) -> List[Dict[str, Any]]:
    suffix = path.suffix.lower()
    batch_id = uuid.uuid4().hex[:12]
    candidates: List[Dict[str, Any]] = []
    try:
        if suffix in {".xlsx", ".xls"}:
            import pandas as pd
            xl = pd.ExcelFile(path)
            for sheet in xl.sheet_names:
                # First try pandas normal header behavior.
                df_header = pd.read_excel(path, sheet_name=sheet, dtype=str).fillna("")
                _dataframe_candidates_from_sheet(df_header, f"{original_name}/{sheet}", batch_id, candidates, max_rows)
                if len(candidates) >= max_rows:
                    break
                # If no candidates from this sheet, try header=None to locate actual header rows.
                if not any(str(c.get("source_file", "")).endswith(f"/{sheet}") for c in candidates):
                    df_raw = pd.read_excel(path, sheet_name=sheet, header=None, dtype=str).fillna("")
                    _dataframe_candidates_from_sheet(df_raw, f"{original_name}/{sheet}", batch_id, candidates, max_rows)
                if len(candidates) >= max_rows:
                    break
        elif suffix == ".csv":
            import pandas as pd
            try:
                df = pd.read_csv(path, dtype=str).fillna("")
            except Exception:
                df = pd.read_csv(path, dtype=str, header=None).fillna("")
            _dataframe_candidates_from_sheet(df, original_name, batch_id, candidates, max_rows)
        elif suffix in {".docx", ".pdf", ".txt", ".md"}:
            parsed = parse_case_file(path)
            numbered_cases = parsed.get("cases") or []
            if numbered_cases:
                for fields in numbered_cases[: max_rows - len(candidates)]:
                    fields = dict(fields)
                    fields.update({"batch_id": batch_id, "candidate_id": f"{batch_id}-{len(candidates)+1:04d}", "row_index": len(candidates), "source_file": original_name})
                    fields["display_title"] = "｜".join([x for x in [fields.get("source_case_number") and f"原编号{fields.get('source_case_number')}", fields.get("patient_name"), fields.get("diagnosis") or fields.get("pathology")] if x])[:120] or Path(original_name).stem
                    candidates.append(_json_safe(fields))
            elif parsed.get("ok") and parsed.get("fields"):
                fields = parsed.get("fields") or {}
                fields.update({"batch_id": batch_id, "candidate_id": f"{batch_id}-0001", "row_index": 0, "source_file": original_name})
                fields["display_title"] = "｜".join([x for x in [fields.get("age") and str(fields.get("age"))+"岁", fields.get("sex"), fields.get("diagnosis") or fields.get("pathology")] if x])[:120] or Path(original_name).stem
                candidates.append(_json_safe(fields))
    except Exception:
        return []
    # Deduplicate by visible title/free text to avoid the same row being parsed twice through two header strategies.
    unique: List[Dict[str, Any]] = []
    seen = set()
    for c in candidates:
        key = re.sub(r"\s+", "", str(c.get("display_title") or "") + str(c.get("free_text") or ""))[:600]
        if key and key not in seen:
            seen.add(key)
            c["row_index"] = len(unique)
            c["candidate_id"] = f"{batch_id}-{len(unique)+1:04d}"
            c["candidate_url"] = f"/candidate/{batch_id}/{c['row_index']}"
            unique.append(c)
    if unique:
        _atomic_write_json(CANDIDATE_BATCH_DIR / f"{batch_id}.json", unique)
    return unique

def _load_candidate_batch(batch_id: str) -> List[Dict[str, Any]]:
    path = CANDIDATE_BATCH_DIR / f"{batch_id}.json"
    if not path.exists():
        return []
    try:
        rows = json.loads(path.read_text(encoding="utf-8"))
        return rows if isinstance(rows, list) else []
    except Exception:
        return []


def _candidate_text(c: Dict[str, Any]) -> str:
    return "\n".join(str(c.get(k) or "") for k in ["name", "sex", "age", "diagnosis", "ls", "tnm", "lymph_node", "history", "symptoms", "tumor", "imaging", "pathology", "surgery", "other_treatment", "followup", "free_text"])


def _find_candidate_matches(attachments: List[Dict[str, Any]], query: str, limit: int = 4) -> List[Dict[str, Any]]:
    batches = [a.get("batch_id") for a in attachments or [] if isinstance(a, dict) and a.get("type") == "candidate_case_batch"]
    if not batches:
        return []
    from core.risk_engine import _char_ngrams
    qgrams = _char_ngrams(query)
    scored=[]
    for bid in batches:
        for c in _load_candidate_batch(str(bid)):
            t = _candidate_text(c)
            grams = _char_ngrams(t)
            base = len(qgrams & grams) / len(qgrams | grams) if (qgrams or grams) else 0
            bonus = 0
            for key in ["name", "age", "sex", "diagnosis", "tnm"]:
                v = str(c.get(key) or "").strip().lower()
                if v and v in query.lower(): bonus += 0.08
            score = base + bonus
            if score > 0.01:
                item=dict(c); item["match_score"] = round(score,3); item["candidate_url"] = f"/candidate/{bid}/{item.get('row_index',0)}"
                scored.append(item)
    scored.sort(key=lambda x:x.get("match_score",0), reverse=True)
    return _json_safe(scored[:limit])

def _load_user_cases() -> None:
    rows = _read_json_with_backup(USER_CASES_PATH, [])
    if not isinstance(rows, list):
        rows = []
    existing = {r.case_id for r in kb.records}
    for row in rows:
        case_id = str(row.get("case_id") or "").strip()
        if not case_id or case_id in existing:
            continue
        try:
            rec = CaseRecord(
                case_id=case_id,
                sheet=str(row.get("sheet") or "用户新增"),
                diagnosis=str(row.get("diagnosis") or ""),
                sex=str(row.get("sex") or ""),
                age=float(row.get("age")) if str(row.get("age") or "").strip() else None,
                history=str(row.get("history") or row.get("free_text") or ""),
                prior_operation=str(row.get("prior_operation") or ""),
                symptoms=str(row.get("symptoms") or ""),
                tumor=str(row.get("tumor") or ""),
                grade=str(row.get("grade") or ""),
                tnm=str(row.get("tnm") or ""),
                lymph_node=str(row.get("lymph_node") or ""),
                ls=str(row.get("ls") or ""),
                immuno=str(row.get("immuno") or ""),
                imaging=str(row.get("imaging") or ""),
                pathology=str(row.get("pathology") or ""),
                surgery=str(row.get("surgery") or ""),
                other_treatment=str(row.get("other_treatment") or ""),
                recurrence=str(row.get("recurrence") or ""),
                followup=str(row.get("followup") or ""),
                remarks=str(row.get("remarks") or ""),
                source_row=int(row.get("source_row") or 0),
                medical_images=_normalize_medical_images(row.get("medical_images") or []),
                patient_name=str(row.get("patient_name") or ""),
                source_case_number=_optional_int(row.get("source_case_number")),
                source_document=str(row.get("source_document") or ""),
                source_pages=_normalize_source_pages(row.get("source_pages") or []),
                source_hash=str(row.get("source_hash") or ""),
                source_import_key=str(row.get("source_import_key") or ""),
                imaging_findings=[dict(x) for x in (row.get("imaging_findings") or []) if isinstance(x, dict)],
            )
            rec.case_signature = str(row.get("case_signature") or "")
            kb.records.append(rec)
            existing.add(case_id)
            if isinstance(row.get("tags"), list):
                tags_map = _load_case_tags()
                if case_id not in tags_map:
                    tags_map[case_id] = [str(x).strip() for x in row.get("tags") if str(x).strip()]
                    _save_case_tags(tags_map)
        except Exception:
            continue


def _save_user_cases() -> None:
    """Persist all doctor-fed cases regardless of editable label.

    A case is considered doctor-fed by its stable USER-* id or import remark,
    not by its display label. This prevents data loss when doctors rename the
    label from the default “用户新增” to another custom group.
    """
    rows = []
    now = datetime.now().isoformat(timespec="seconds")
    for r in kb.records:
        if _is_user_case(r):
            r.case_signature = _case_duplicate_fingerprint(_case_fields_from_record(r))
            d = _record_to_saved_dict(r)
            d["persisted_at"] = now
            rows.append(d)
    _atomic_write_json(USER_CASES_PATH, rows)
    active_ids = {str(row.get("case_id") or "") for row in rows}
    deleted = _load_deleted_case_ids()
    if active_ids & deleted:
        _atomic_write_json(DELETED_CASES_PATH, sorted(deleted - active_ids))
    # Read-back verification marker. If the file cannot be read, the manifest
    # will expose the mismatch instead of silently losing data.
    try:
        persisted_count = len(json.loads(USER_CASES_PATH.read_text(encoding="utf-8")))
    except Exception:
        persisted_count = -1
    _save_library_snapshot()
    _write_storage_manifest()


def _next_user_case_id() -> str:
    nums = []
    known_ids = [r.case_id for r in kb.records] + list(_load_deleted_case_ids())
    for case_id in known_ids:
        if case_id.startswith("USER-"):
            try:
                nums.append(int(case_id.split("-", 1)[1]))
            except Exception:
                pass
    return f"USER-{(max(nums) if nums else 0) + 1:03d}"



CASE_DUPLICATE_FIELDS = [
    "diagnosis", "sex", "age", "history", "prior_operation", "symptoms", "tumor",
    "grade", "tnm", "lymph_node", "ls", "immuno", "imaging", "pathology",
    "surgery", "other_treatment", "recurrence", "followup",
]


def _normalize_case_duplicate_value(value: Any) -> str:
    text = str(value or "").strip().lower()
    if text in {"", "nan", "none", "null", "/", "\\", "未记录", "未填写", "暂无", "待整理病例"}:
        return ""
    if re.fullmatch(r"\d+\.0+", text):
        text = text.split(".", 1)[0]
    return re.sub(r"[\s,，。；;：:、|/\\()（）\[\]{}<>\-_]+", "", text)


def _case_duplicate_payload(fields: Dict[str, Any], source_text: str = "") -> Dict[str, str]:
    fields = fields if isinstance(fields, dict) else {}
    raw = {
        "diagnosis": fields.get("diagnosis") or fields.get("主要诊断") or "",
        "sex": fields.get("sex") or fields.get("性别") or "",
        "age": fields.get("age") or fields.get("年龄") or "",
        "history": fields.get("history") or fields.get("free_text") or source_text or "",
        "prior_operation": fields.get("prior_operation") or "",
        "symptoms": fields.get("symptoms") or "",
        "tumor": fields.get("tumor") or fields.get("tumor_location") or "",
        "grade": fields.get("grade") or "",
        "tnm": fields.get("tnm") or "",
        "lymph_node": fields.get("lymph_node") or "",
        "ls": fields.get("ls") or "",
        "immuno": fields.get("immuno") or "",
        "imaging": fields.get("imaging") or "",
        "pathology": fields.get("pathology") or fields.get("病理") or "",
        "surgery": fields.get("surgery") or "",
        "other_treatment": fields.get("other_treatment") or "",
        "recurrence": fields.get("recurrence") or "",
        "followup": fields.get("followup") or "",
    }
    return {key: val for key, value in raw.items() if (val := _normalize_case_duplicate_value(value))}


def _case_duplicate_fingerprint(fields: Dict[str, Any], source_text: str = "") -> str:
    payload = _case_duplicate_payload(fields, source_text)
    # Diagnosis-only rows are not enough to prove that two patients are identical.
    if len(payload) < 3 and sum(len(v) for v in payload.values()) < 60:
        return ""
    normalized = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(normalized.encode("utf-8", errors="ignore")).hexdigest()


def _case_fields_from_record(rec: CaseRecord) -> Dict[str, Any]:
    return {key: getattr(rec, key, "") for key in CASE_DUPLICATE_FIELDS}


def _case_signature_fields(fields: Dict[str, Any], source_text: str = "") -> str:
    """Create a restart-safe signature from fields that are actually persisted."""
    return _case_duplicate_fingerprint(fields, source_text)


def _find_duplicate_case(fields: Dict[str, Any], source_text: str = "") -> CaseRecord | None:
    fingerprint = _case_duplicate_fingerprint(fields, source_text)
    if not fingerprint:
        return None
    for rec in kb.records:
        if not _is_user_case(rec):
            continue
        if _case_duplicate_fingerprint(_case_fields_from_record(rec)) == fingerprint:
            return rec
    return None

def _add_case_from_fields(fields: Dict[str, Any], source_text: str = "") -> CaseRecord:
    free_text = str(fields.get("free_text") or source_text or "").strip()
    rec = CaseRecord(
        case_id=_next_user_case_id(),
        sheet=_sanitize_case_label(fields.get("case_label") or fields.get("sheet") or fields.get("分组标签") or fields.get("标签")),
        diagnosis=str(fields.get("diagnosis") or fields.get("主要诊断") or "待整理病例"),
        sex=str(fields.get("sex") or ""),
        age=float(fields.get("age")) if str(fields.get("age") or "").strip().replace('.', '', 1).isdigit() else None,
        history=str(fields.get("history") or free_text),
        prior_operation=str(fields.get("prior_operation") or ""),
        symptoms=str(fields.get("symptoms") or ""),
        tumor=str(fields.get("tumor") or fields.get("tumor_location") or ""),
        grade=str(fields.get("grade") or ""),
        tnm=str(fields.get("tnm") or ""),
        lymph_node=str(fields.get("lymph_node") or ""),
        ls=str(fields.get("ls") or ""),
        immuno=str(fields.get("immuno") or ""),
        imaging=str(fields.get("imaging") or ""),
        pathology=str(fields.get("pathology") or ""),
        surgery=str(fields.get("surgery") or ""),
        other_treatment=str(fields.get("other_treatment") or ""),
        recurrence=str(fields.get("recurrence") or ""),
        followup=str(fields.get("followup") or ""),
        remarks=str(fields.get("remarks") or "由对话或上传文件加入。"),
        source_row=0,
        medical_images=_normalize_medical_images(fields.get("medical_images") or []),
        patient_name=str(fields.get("patient_name") or fields.get("name") or fields.get("姓名") or ""),
        source_case_number=_optional_int(fields.get("source_case_number")),
        source_document=str(fields.get("source_document") or ""),
        source_pages=_normalize_source_pages(fields.get("source_pages") or []),
        source_hash=str(fields.get("source_hash") or ""),
        source_import_key=str(fields.get("source_import_key") or ""),
        imaging_findings=[dict(x) for x in (fields.get("imaging_findings") or []) if isinstance(x, dict)],
    )
    rec.case_signature = _case_signature_fields(fields, source_text)
    kb.records.append(rec)
    _save_user_cases()
    return rec


def _find_case(case_id: str):
    for rec in kb.records:
        if rec.case_id == case_id:
            return _case_to_public_dict(rec)
    return None


_ensure_storage_files()
_load_user_cases()
_normalize_legacy_user_labels()
_apply_deleted_case_filter()


@app.after_request
def add_no_cache_headers(response):
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response


@app.errorhandler(Exception)
def handle_exception(exc):
    if request.path.startswith("/api/"):
        return jsonify({"ok": False, "error": str(exc), "traceback": traceback.format_exc(limit=2)}), 500
    raise exc


@app.get("/")
def index():
    return render_template("index.html")


@app.get("/favicon.ico")
def favicon():
    return send_from_directory(BASE_DIR / "static" / "assets", "app_icon.ico", mimetype="image/vnd.microsoft.icon")


@app.get("/knowledge")
def knowledge():
    return render_template("knowledge.html")


@app.get("/add-case")
def add_case_page():
    return render_template("add_case.html")


@app.get("/articles")
def articles_page():
    return render_template("articles.html")


@app.get("/article/<article_id>")
def article_detail(article_id):
    article = next((a for a in _load_articles() if str(a.get("article_id")) == article_id), None)
    if not article:
        abort(404)
    return render_template("article_detail.html", article=article)


@app.get("/uploads/<path:filename>")
def uploaded_file(filename):
    return send_from_directory(UPLOAD_DIR, filename)


@app.get("/case/<case_id>")
def case_detail(case_id):
    case = _find_case(case_id)
    if not case:
        abort(404)
    fields = [
        ("性别", "sex"), ("年龄", "age"), ("LS/硬化性苔藓", "ls"),
        ("病史要点", "history"), ("既往尿道操作/重建史", "prior_operation"),
        ("症状/体征", "symptoms"), ("肿瘤情况", "tumor"),
        ("分化程度", "grade"), ("TNM/分期", "tnm"), ("淋巴结状态", "lymph_node"),
        ("影像/尿道镜", "imaging"), ("病理", "pathology"),
        ("免疫组化", "immuno"), ("手术治疗", "surgery"), ("其他治疗/用药", "other_treatment"),
        ("复发/转移", "recurrence"),
        ("随访", "followup"), ("备注", "remarks"),
    ]
    return render_template("case_detail.html", case=case, fields=fields)


@app.get("/api/case/<case_id>")
def api_case_detail(case_id):
    case = _find_case(case_id)
    if not case:
        return jsonify({"ok": False, "error": "未找到该病例"}), 404
    return jsonify({"ok": True, "case": case})


@app.get("/api/summary")
def api_summary():
    summary = kb.summary()
    summary["articles"] = len(_load_articles())
    return jsonify(summary)


@app.get("/api/cases")
def api_cases():
    sheet = request.args.get("sheet") or None
    q = (request.args.get("q") or "").strip()
    try:
        limit = int(request.args.get("limit", 200))
    except ValueError:
        limit = 200
    records = [r for r in kb.records if sheet is None or r.sheet == sheet]
    if q:
        records = [r for r in records if _case_matches_query(r, q)]
    records = sorted(records, key=case_sort_key)
    cases = [_case_to_public_dict(r) for r in records[:limit]]
    return jsonify({"ok": True, "cases": cases, "total": len(records), "query": q, "sheet": sheet or ""})



@app.get("/api/articles")
def api_articles():
    q = (request.args.get("q") or "").strip()
    try:
        limit = int(request.args.get("limit", 100))
    except ValueError:
        limit = 100
    rows = [a for a in _load_articles() if _article_matches_query(a, q)]
    rows = rows[:limit]
    for item in rows:
        content = str(item.get("content") or "")
        item["content_preview"] = content[:360] + ("..." if len(content) > 360 else "")
        item["article_url"] = f"/article/{item.get('article_id')}"
    return jsonify({"ok": True, "articles": _json_safe(rows), "total": len(rows), "query": q})




@app.delete("/api/article/<article_id>")
def api_delete_article(article_id):
    rows = _load_articles()
    before = len(rows)
    rows = [a for a in rows if str(a.get("article_id")) != str(article_id)]
    if len(rows) == before:
        return jsonify({"ok": False, "error": "未找到该文章"}), 404
    _save_articles(rows)
    return jsonify({"ok": True, "deleted": article_id, "before": before, "after": len(rows), "message": f"已删除文章：{article_id}"})


@app.delete("/api/articles/bulk")
def api_delete_articles_bulk():
    payload = request.get_json(force=True, silent=True) or {}
    article_ids = [str(x).strip() for x in (payload.get("article_ids") or []) if str(x).strip()]
    confirm = str(payload.get("confirm") or "")
    if confirm != "DELETE":
        return jsonify({"ok": False, "error": "缺少确认字段。"}), 400
    if not article_ids:
        return jsonify({"ok": False, "error": "没有选择要删除的文章。"}), 400
    id_set = set(article_ids)
    rows = _load_articles()
    before = len(rows)
    rows = [a for a in rows if str(a.get("article_id")) not in id_set]
    _save_articles(rows)
    return jsonify({"ok": True, "deleted": sorted(id_set), "before": before, "after": len(rows), "message": f"已删除 {before - len(rows)} 篇文章。"})


@app.get("/api/storage/status")
def api_storage_status():
    return jsonify({
        "ok": True,
        "data_dir": str(PERSISTENT_DATA_DIR.resolve()),
        "project_data_dir": str(PROJECT_DATA_DIR.resolve()),
        "user_cases_path": str(USER_CASES_PATH),
        "articles_path": str(ARTICLES_PATH),
        "case_tags_path": str(CASE_TAGS_PATH),
        "user_cases_backup_path": str(USER_CASES_PATH.with_suffix(USER_CASES_PATH.suffix + ".bak")),
        "storage_manifest_path": str(PERSISTENT_DATA_DIR / "storage_manifest.json"),
        "library_state_path": str(LIBRARY_STATE_PATH),
        "knowledge_digest_path": str(KB_DIGEST_PATH),
        "persistence_repair_log_path": str(PERSISTENCE_REPAIR_LOG_PATH),
        "migration_state_path": str(MIGRATION_STATE_PATH),
        "persistent_data_dir": str(PERSISTENT_DATA_DIR),
        "uploads_dir": str(UPLOAD_DIR),
        "upload_files": len([p for p in UPLOAD_DIR.iterdir() if p.is_file()]),
        "user_cases_saved": USER_CASES_PATH.exists(),
        "articles_saved": ARTICLES_PATH.exists(),
        "user_cases": len([r for r in kb.records if _is_user_case(r)]),
        "articles": len(_load_articles()),
        "api_config_saved": API_CONFIG_PATH.exists(),
        "api_config_history_count": len(_normalize_api_config_history(save=False)),
        "articles_deleted_marker": ARTICLES_DELETED_MARKER.exists(),
        "articles_v34_cleared_marker": ARTICLES_V34_CLEARED_MARKER.exists(),
        "max_upload_mb": app.config.get("MAX_CONTENT_LENGTH", 0) // 1024 // 1024,
        "max_batch_files": MAX_BATCH_FILES,
    })

@app.post("/api/articles")
def api_add_article():
    payload = request.get_json(force=True, silent=True) or {}
    fields = payload.get("fields") or payload
    try:
        article = _add_article(fields)
    except ValueError as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400
    if article.get("duplicate"):
        return jsonify({"ok": True, "duplicate": True, "article": _json_safe(article), "message": article.get("message")})
    return jsonify({"ok": True, "article": _json_safe(article), "message": f"已加入文献库：{article['article_id']}"})


@app.post("/api/articles/autosave")
def api_articles_autosave():
    payload = request.get_json(force=True, silent=True) or {}
    fields = payload.get("fields") or {}
    article_id = str(payload.get("article_id") or "").strip()
    rows = _load_articles()
    now = datetime.now().isoformat(timespec="seconds")
    if article_id:
        for idx, row in enumerate(rows):
            if str(row.get("article_id")) == article_id:
                merged = {**row, **fields, "article_id": article_id, "updated_at": now}
                if not str(merged.get("title") or "").strip():
                    merged["title"] = "未命名文献"
                dup = _find_duplicate_article(merged, exclude_id=article_id)
                if dup:
                    return jsonify({"ok": True, "duplicate": True, "article": _json_safe(dup), "message": f"检测到相同或高度重复文章：{dup.get('article_id')}，未重复保存。"})
                merged["signature"] = _article_signature_fields(merged)
                rows[idx] = _json_safe(merged)
                _save_articles(rows)
                return jsonify({"ok": True, "article": _json_safe(merged), "message": f"已自动更新：{article_id}"})
    article = _add_article({**fields, "created_at": now, "updated_at": now})
    if article.get("duplicate"):
        return jsonify({"ok": True, "duplicate": True, "article": _json_safe(article), "message": article.get("message")})
    return jsonify({"ok": True, "article": _json_safe(article), "message": f"已自动保存：{article['article_id']}"})


@app.post("/api/upload_article")
def api_upload_article():
    if "file" not in request.files:
        return jsonify({"ok": False, "error": "没有收到文件字段 file"}), 400
    file = request.files["file"]
    if not file.filename:
        return jsonify({"ok": False, "error": "文件名为空"}), 400
    original_name = file.filename
    suffix = Path(original_name).suffix.lower()
    if suffix not in ARTICLE_EXTENSIONS:
        return jsonify({"ok": False, "error": "文章投喂支持 txt/docx/pdf/xlsx/xls/csv/md。"}), 400
    safe_stem = Path(secure_filename(original_name)).stem or "article"
    stored_name = f"{uuid.uuid4().hex}_{safe_stem}{suffix}"
    dest = UPLOAD_DIR / stored_name
    file.save(dest)
    text = _extract_article_text(dest)
    images = _extract_article_images(dest)
    saved_api = _load_saved_api_config()
    image_analysis = _summarize_article_images_for_doctor(
        images,
        original_name,
        article_text=text,
        api_key=request.form.get("api_key", "") or saved_api.get("api_key", ""),
        provider=request.form.get("provider", "") or saved_api.get("provider", ""),
        model=request.form.get("model", "") or saved_api.get("model", ""),
        base_url=request.form.get("base_url", "") or saved_api.get("base_url", ""),
    )
    full_text = (text + image_analysis)[:MAX_TEXT_CHARS_PER_FILE]
    note = f"已解析文章文本 {len(text)} 字符。"
    if images:
        note += f" 检测到 {len(images)} 张图片，已保存并尝试分析。"
    return jsonify({
        "ok": True,
        "filename": original_name,
        "stored_as": stored_name,
        "url": f"/uploads/{stored_name}",
        "text": full_text,
        "images": _json_safe(images),
        "note": note + " 可编辑后加入文献库。"
    })


@app.post("/api/upload_articles_batch")
def api_upload_articles_batch():
    files = request.files.getlist("files")
    if not files:
        return jsonify({"ok": False, "error": "没有收到文件字段 files"}), 400
    if len(files) > MAX_BATCH_FILES:
        return jsonify({"ok": False, "error": f"单次最多处理 {MAX_BATCH_FILES} 个文件；请分批导入。"}), 413
    added, failed = [], []
    for file in files:
        if not file or not file.filename:
            continue
        original_name = file.filename
        suffix = Path(original_name).suffix.lower()
        if suffix not in ARTICLE_EXTENSIONS:
            failed.append({"filename": original_name, "error": "不支持的文章文件类型"})
            continue
        safe_stem = Path(secure_filename(original_name)).stem or "article"
        stored_name = f"{uuid.uuid4().hex}_{safe_stem}{suffix}"
        dest = UPLOAD_DIR / stored_name
        file.save(dest)
        try:
            text = _extract_article_text(dest)
            images = _extract_article_images(dest)
            saved_api = _load_saved_api_config()
            image_analysis = _summarize_article_images_for_doctor(
                images,
                original_name,
                article_text=text,
                api_key=request.form.get("api_key", "") or saved_api.get("api_key", ""),
                provider=request.form.get("provider", "") or saved_api.get("provider", ""),
                model=request.form.get("model", "") or saved_api.get("model", ""),
                base_url=request.form.get("base_url", "") or saved_api.get("base_url", ""),
            )
            article = _add_article({
                "title": Path(original_name).stem,
                "content": (text + image_analysis)[:MAX_TEXT_CHARS_PER_FILE],
                "source_file": original_name,
                "source_url": f"/uploads/{stored_name}",
                "stored_as": stored_name,
                "article_images": images,
                "created_at": datetime.now().isoformat(timespec="seconds"),
            })
            if article.get("duplicate"):
                failed.append({"filename": original_name, "error": article.get("message") or "重复文章，已跳过"})
            else:
                added.append(_json_safe(article))
        except Exception as exc:
            failed.append({"filename": original_name, "error": str(exc)})
    return jsonify({"ok": True, "added": added, "failed": failed, "message": f"已加入 {len(added)} 篇文章，失败 {len(failed)} 个文件。"})

@app.post("/api/analyze")
def api_analyze():
    payload = request.get_json(force=True, silent=True) or {}
    report = generate_traceable_report(kb, payload, top_n=4)
    return jsonify({"ok": True, **report})


def _has_nonempty_case_content(fields: Dict[str, Any], source_text: str = "") -> bool:
    """Prevent saving blank or placeholder-only cases."""
    if str(source_text or "").strip():
        return True
    meaningful_keys = [
        "diagnosis", "主要诊断", "sex", "age", "ls", "tnm", "grade", "lymph_node",
        "history", "prior_operation", "symptoms", "tumor", "tumor_location", "imaging",
        "pathology", "immuno", "surgery", "other_treatment", "recurrence", "followup",
        "remarks", "free_text"
    ]
    placeholder_values = {"", "nan", "none", "null", "/", "\\", "待整理病例", "未填写", "暂无"}
    for key in meaningful_keys:
        val = fields.get(key) if isinstance(fields, dict) else ""
        if val is None:
            continue
        text = str(val).strip()
        if text and text.lower() not in placeholder_values:
            return True
    images = fields.get("medical_images") if isinstance(fields, dict) else None
    return bool(images)


@app.post("/api/add_case")
def api_add_case():
    payload = request.get_json(force=True, silent=True) or {}
    fields = payload.get("fields") or {}
    source_text = payload.get("source_text") or ""
    if not _has_nonempty_case_content(fields, source_text):
        return jsonify({"ok": False, "error": "不能保存空白病例。请填写病例摘要，或上传并解析病例文件。"}), 400
    duplicate = _find_duplicate_case(fields, source_text)
    if duplicate:
        return jsonify({"ok": True, "duplicate": True, "case": _case_to_public_dict(duplicate), "message": f"检测到相同或高度相似病例：{duplicate.case_id}，已跳过重复保存。"})
    rec = _add_case_from_fields(fields, source_text=source_text)
    return jsonify({"ok": True, "case": _case_to_public_dict(rec), "message": f"已加入知识库：{rec.case_id}"})


@app.post("/api/upload_cases_batch")
def api_upload_cases_batch():
    files = request.files.getlist("files")
    if not files:
        return jsonify({"ok": False, "error": "没有收到文件字段 files"}), 400
    if len(files) > MAX_BATCH_FILES:
        return jsonify({"ok": False, "error": f"单次最多处理 {MAX_BATCH_FILES} 个文件；请分批导入。"}), 413
    candidates, failed = [], []
    for file in files:
        if not file or not file.filename:
            continue
        original_name = file.filename
        suffix = Path(original_name).suffix.lower()
        if suffix not in CASE_TABLE_EXTENSIONS:
            failed.append({"filename": original_name, "error": "不支持的病例文件类型"})
            continue
        safe_stem = Path(secure_filename(original_name)).stem or "case"
        stored_name = f"{uuid.uuid4().hex}_{safe_stem}{suffix}"
        dest = UPLOAD_DIR / stored_name
        file.save(dest)
        try:
            extracted = _extract_case_batch_candidates(dest, original_name) if suffix in CASE_TABLE_EXTENSIONS else []
            if not extracted:
                parsed = parse_case_file(dest)
                if parsed.get("ok") and parsed.get("fields"):
                    fields = parsed.get("fields") or {}
                    fields["source_file"] = original_name
                    fields["stored_as"] = stored_name
                    fields["source_url"] = f"/uploads/{stored_name}"
                    fields["display_title"] = "｜".join([x for x in [fields.get("age") and str(fields.get("age"))+"岁", fields.get("sex"), fields.get("diagnosis") or fields.get("pathology")] if x])[:120] or Path(original_name).stem
                    extracted = [fields]
            for c in extracted:
                c = dict(c)
                c["source_url"] = c.get("source_url") or f"/uploads/{stored_name}"
                c["stored_as"] = c.get("stored_as") or stored_name
                candidates.append(_json_safe(c))
        except Exception as exc:
            failed.append({"filename": original_name, "error": str(exc)})
    return jsonify({"ok": True, "candidates": candidates[:2000], "failed": failed, "message": f"已识别 {len(candidates)} 个候选病例，失败 {len(failed)} 个文件。请在弹窗中选择需要加入知识库的病例。"})

@app.post("/api/candidates/add_many")
def api_candidates_add_many():
    payload = request.get_json(force=True, silent=True) or {}
    candidates = payload.get("candidates") or []
    if not isinstance(candidates, list) or not candidates:
        return jsonify({"ok": False, "error": "未选择候选病例。"}), 400
    added=[]
    skipped=[]
    for cand in candidates[:5000]:
        if isinstance(cand, dict):
            source_text = str(cand.get("free_text") or cand.get("history") or "")
            duplicate = _find_duplicate_case(cand, source_text)
            if duplicate:
                skipped.append(duplicate.case_id)
                continue
            rec = _add_case_from_fields(cand, source_text=source_text)
            added.append(_case_to_public_dict(rec))
    msg = f"已加入 {len(added)} 个病例" + (f"，跳过重复 {len(skipped)} 个。" if skipped else "。")
    return jsonify({"ok": True, "added": added, "skipped_duplicates": skipped, "message": msg})


@app.delete("/api/cases/bulk")
def api_delete_cases_bulk():
    payload = request.get_json(force=True, silent=True) or {}
    case_ids = [str(x).strip() for x in (payload.get("case_ids") or []) if str(x).strip()]
    confirm = str(payload.get("confirm") or "")
    if confirm != "DELETE":
        return jsonify({"ok": False, "error": "缺少确认字段。"}), 400
    if not case_ids:
        return jsonify({"ok": False, "error": "没有选择要删除的病例。"}), 400
    existing = {r.case_id for r in kb.records}
    to_delete = [cid for cid in case_ids if cid in existing]
    deleted = _load_deleted_case_ids()
    for cid in to_delete:
        deleted.add(cid)
    before = len(kb.records)
    kb.records = [r for r in kb.records if r.case_id not in set(to_delete)]
    _save_user_cases()
    _save_deleted_case_ids(deleted)
    return jsonify({"ok": True, "deleted": to_delete, "before": before, "after": len(kb.records), "message": f"已删除 {len(to_delete)} 个病例。"})



@app.patch("/api/case/<case_id>/tags")
def api_update_case_tags(case_id):
    payload = request.get_json(force=True, silent=True) or {}
    tags_raw = payload.get("tags", [])
    if isinstance(tags_raw, str):
        tags = [x.strip() for x in re.split(r"[,，;；\s]+", tags_raw) if x.strip()]
    elif isinstance(tags_raw, list):
        tags = [str(x).strip() for x in tags_raw if str(x).strip()]
    else:
        return jsonify({"ok": False, "error": "标签格式错误。"}), 400
    tags = list(dict.fromkeys(tags))[:12]
    target = next((r for r in kb.records if r.case_id == case_id), None)
    if not target:
        return jsonify({"ok": False, "error": "未找到该病例。"}), 404
    rows = _load_case_tags()
    rows[case_id] = tags
    _save_case_tags(rows)
    if _is_user_case(target):
        _save_user_cases()
    return jsonify({"ok": True, "case_id": case_id, "tags": tags, "message": "标签已更新。"})

@app.patch("/api/case/<case_id>/label")
def api_update_case_label(case_id):
    payload = request.get_json(force=True, silent=True) or {}
    label = _sanitize_case_label(payload.get("label") or payload.get("sheet") or payload.get("case_label"))
    target = next((r for r in kb.records if r.case_id == case_id), None)
    if not target:
        return jsonify({"ok": False, "error": "未找到该病例。"}), 404
    if not _is_user_case(target):
        return jsonify({"ok": False, "error": "默认 Excel 数据库中的原始分组不建议修改；请仅修改用户新增的标签。"}), 400
    target.sheet = label
    _save_user_cases()
    return jsonify({"ok": True, "case_id": case_id, "label": label, "message": "病例标签已更新。"})


@app.delete("/api/case/<case_id>")
def api_delete_case(case_id):
    before = len(kb.records)
    target = None
    for r in kb.records:
        if r.case_id == case_id:
            target = r
            break
    if not target:
        return jsonify({"ok": False, "error": "未找到该病例"}), 404
    kb.records = [r for r in kb.records if r.case_id != case_id]
    deleted = _load_deleted_case_ids()
    deleted.add(case_id)
    if _is_user_case(target):
        _save_user_cases()
    _save_deleted_case_ids(deleted)
    return jsonify({"ok": True, "deleted": case_id, "before": before, "after": len(kb.records)})




def _is_analysis_request(question: str, mode: str = "") -> bool:
    """Only the first confirmed-patient analysis and explicit requests should pull case/article cards.
    Normal follow-up questions are answered by the AI using current patient context, without appending cases/articles.
    """
    q = str(question or "")
    if mode == "initial_patient_analysis":
        return True
    keywords = [
        "总体分析", "分析当前", "完整分析", "治疗讨论", "治疗建议", "手术", "用药",
        "相似病例", "参考文献", "文献", "指南", "证据", "转归", "复发", "预后", "方案", "下一步"
    ]
    return any(k in q for k in keywords)




@app.get("/api/settings/mascot")
def api_get_mascot():
    return jsonify({
        "ok": True,
        "stored_as": "mascot_fixed.png",
        "url": "/static/assets/mascot_fixed.png",
        "updated_at": datetime.now().isoformat(timespec="seconds"),
    })


@app.post("/api/settings/mascot")
def api_set_mascot():
    return jsonify({"ok": False, "error": "当前版本已固定使用项目头像，不再支持上传更换。"}), 400


def _load_saved_api_config() -> Dict[str, Any]:
    try:
        data = json.loads(API_CONFIG_PATH.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _load_api_config_history() -> List[Dict[str, Any]]:
    try:
        data = json.loads(API_CONFIG_HISTORY_PATH.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _normalize_api_config_history(save: bool = True) -> List[Dict[str, Any]]:
    """Keep only configs that were successfully tested and ensure every item has config_id.

    v33 stored some failed / incomplete entries. v34 prunes them so the
    selection panel only contains callable saved configurations.
    """
    rows = []
    seen: set[str] = set()
    for item in _load_api_config_history():
        if not isinstance(item, dict):
            continue
        clean = {k: str(item.get(k) or "").strip() for k in ["api_key", "provider", "model", "base_url", "updated_at", "config_id"]}
        if not clean.get("api_key") or not clean.get("provider") or not clean.get("model"):
            continue
        # Delete API configs that were never verified successfully. Legacy rows
        # without a flag are treated as unverified unless they are the current
        # saved config and get re-saved successfully later.
        if item.get("test_ok") is not True:
            continue
        if not clean.get("config_id"):
            clean["config_id"] = _api_config_id(clean)
        clean["test_ok"] = True
        if clean["config_id"] in seen:
            continue
        seen.add(clean["config_id"])
        rows.append(clean)
    if save:
        _atomic_write_json(API_CONFIG_HISTORY_PATH, rows[:12])
    return rows[:12]


def _prune_current_api_config_if_unverified() -> None:
    cfg = _load_saved_api_config()
    if not cfg:
        return
    if cfg.get("test_ok") is not True:
        try:
            API_CONFIG_PATH.unlink(missing_ok=True)
        except Exception:
            pass


def _mask_api_key(key: str) -> str:
    key = str(key or "")
    return key[:4] + "..." + key[-4:] if len(key) > 8 else ("已保存" if key else "")


def _api_config_id(data: Dict[str, Any]) -> str:
    compact = "|".join(str(data.get(k) or "") for k in ["provider", "model", "base_url", "api_key"])
    return hashlib.sha1(compact.encode("utf-8", errors="ignore")).hexdigest()[:16]


def _masked_api_config(data: Dict[str, Any]) -> Dict[str, Any]:
    out = {k: data.get(k, "") for k in ["provider", "model", "base_url", "updated_at", "config_id"]}
    out["api_key_masked"] = _mask_api_key(str(data.get("api_key") or ""))
    return out


def _save_api_config(payload: Dict[str, Any], test_ok: bool = False) -> Dict[str, Any]:
    allowed = {"api_key", "provider", "model", "base_url"}
    data = {k: str(payload.get(k) or "").strip() for k in allowed}
    data["updated_at"] = datetime.now().isoformat(timespec="seconds")
    data["config_id"] = _api_config_id(data)
    data["test_ok"] = bool(test_ok)
    if not data.get("api_key"):
        return data
    if not test_ok:
        # Failed configs should not be remembered or overwrite the active config.
        return data
    _atomic_write_json(API_CONFIG_PATH, data)
    history = _normalize_api_config_history(save=False)
    merged = [data] + [x for x in history if x.get("config_id") != data["config_id"]]
    _atomic_write_json(API_CONFIG_HISTORY_PATH, merged[:12])
    return data


@app.get("/api/llm/config")
def api_llm_config_get():
    _prune_current_api_config_if_unverified()
    cfg = _load_saved_api_config()
    if cfg and not cfg.get("config_id"):
        cfg["config_id"] = _api_config_id(cfg)
    if cfg and cfg.get("test_ok") is not True:
        cfg = {}
    masked = _masked_api_config(cfg) if cfg else {}
    raw_history = _normalize_api_config_history(save=True)
    if cfg and cfg.get("api_key") and cfg.get("test_ok") is True:
        raw_history = [cfg] + [x for x in raw_history if x.get("config_id") != cfg.get("config_id")]
    history = [_masked_api_config(x) for x in raw_history]
    return jsonify({"ok": True, "config": masked, "history": history, "remembered": bool(cfg), "message": "已自动清理未连接成功的 API 配置。"})


@app.post("/api/llm/config/use")
def api_llm_config_use():
    payload = request.get_json(force=True, silent=True) or {}
    config_id = str(payload.get("config_id") or "").strip()
    match = None
    for item in _normalize_api_config_history(save=True):
        if str(item.get("config_id") or "") == config_id:
            match = item
            break
    if not match:
        return jsonify({"ok": False, "error": "未找到该历史 API 配置。"}), 404
    match = {k: str(match.get(k) or "").strip() for k in ["api_key", "provider", "model", "base_url", "updated_at", "config_id"]}
    if not match.get("api_key"):
        return jsonify({"ok": False, "error": "该历史配置没有保存 API Key，无法调用。"}), 400
    match["updated_at"] = datetime.now().isoformat(timespec="seconds")
    match["test_ok"] = True
    _atomic_write_json(API_CONFIG_PATH, match)
    return jsonify({"ok": True, "config": _masked_api_config(match), "message": "已切换到该 API 配置。"})



@app.delete("/api/llm/config")
def api_llm_config_clear():
    """Clear backend-saved API config so the app really returns to local mode."""
    try:
        if API_CONFIG_PATH.exists():
            API_CONFIG_PATH.unlink()
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 500
    return jsonify({"ok": True, "message": "已清除后端 API 配置。"})

@app.post("/api/llm/config")
def api_llm_config_save():
    payload = request.get_json(force=True, silent=True) or {}
    draft = {k: str(payload.get(k) or "").strip() for k in ["api_key", "provider", "model", "base_url"]}
    test = test_llm_connection(
        api_key=draft.get("api_key", ""),
        provider=draft.get("provider", "openai"),
        model=draft.get("model", ""),
        base_url=draft.get("base_url", ""),
    ) if draft.get("api_key") else {"ok": False, "error": "未填写 API Key。"}
    cfg = _save_api_config(draft, test_ok=bool(test.get("ok")))
    _normalize_api_config_history(save=True)
    if test.get("ok"):
        return jsonify({"ok": True, "saved": True, "test": test, "config": _masked_api_config(cfg), "message": "连接成功，已保存并加入已记住配置。"})
    return jsonify({"ok": True, "saved": False, "test": test, "message": "连接失败，未保存该 API 配置。"})


@app.post("/api/llm/test")
def api_llm_test():
    payload = request.get_json(force=True, silent=True) or {}
    result = test_llm_connection(
        api_key=str(payload.get("api_key") or ""),
        provider=str(payload.get("provider") or "openai"),
        model=str(payload.get("model") or ""),
        base_url=str(payload.get("base_url") or ""),
    )
    status = 200 if result.get("ok") else 400
    return jsonify(result), status

@app.post("/api/chat")
def api_chat():
    payload = request.get_json(force=True, silent=True) or {}
    patient = payload.get("patient") or {}
    question = (payload.get("question") or "请结合当前病例进行讨论。").strip()
    mode = (payload.get("mode") or "").strip()
    history = payload.get("history") or []
    attachments = payload.get("attachments") or []
    saved_api = _load_saved_api_config()
    if saved_api and saved_api.get("test_ok") is not True:
        saved_api = {}
    api_key = (payload.get("api_key") or saved_api.get("api_key") or "").strip()
    model = (payload.get("model") or saved_api.get("model") or "").strip()
    provider = (payload.get("provider") or saved_api.get("provider") or "openai").strip()
    base_url = (payload.get("base_url") or saved_api.get("base_url") or "").strip()
    if not patient or not any(str(v).strip() for v in patient.values() if v is not None):
        patient = {"free_text": question}
    image_attachments = [a for a in (attachments or []) if isinstance(a, dict) and a.get("type") == "image"]
    if image_attachments:
        patient = {**patient, "medical_images": image_attachments}

    # 对话内投喂：输入“投喂进入数据库/加入数据库/添加病例”时，将当前会话整理为新增病例。
    feed_words = ["投喂进入数据库", "加入数据库", "添加到数据库", "添加病例", "存入知识库", "加入知识库"]
    added_case = None
    added_case_duplicate = False
    if any(w in question for w in feed_words):
        merged_text = "\n".join([str(m.get("content", "")) for m in history if m.get("role") == "user"] + [question])
        patient = {**patient, "free_text": (patient.get("free_text") or "") + "\n" + merged_text}
        added_case = _add_case_from_fields(patient, source_text=merged_text)

    analysis_request = _is_analysis_request(question, mode)
    if analysis_request:
        report = generate_traceable_report(kb, patient, top_n=4)
        article_query = " ".join([question, str(patient.get("free_text") or ""), str(patient.get("diagnosis") or ""), str(patient.get("pathology") or ""), str(patient.get("surgery") or "")])
        report["related_articles"] = _search_articles(article_query, limit=4)
        candidate_query = " ".join([question, str(patient.get("free_text") or ""), str(patient.get("age") or ""), str(patient.get("sex") or ""), str(patient.get("diagnosis") or "")])
        report["candidate_matches"] = _find_candidate_matches(attachments, candidate_query, limit=4)
        report["knowledge_digest"] = _knowledge_digest()
    else:
        report = {
            "similar_cases": [],
            "related_articles": [],
            "candidate_matches": [],
            "treatment_outcomes": {},
            "risk": {"missing_items": []},
            "answer_mode": "normal_followup_no_forced_retrieval",
            "knowledge_digest": _knowledge_digest(),
        }
    llm_result = ask_llm(question=question, report=report, patient=patient, history=history, attachments=attachments, api_key_override=api_key, model_override=model, mode_override=(mode if analysis_request else "normal_followup"), provider_override=provider, base_url_override=base_url)
    if added_case:
        llm_result["answer"] = f"已将当前对话整理并加入知识库：{added_case.case_id}。\n\n" + llm_result.get("answer", "")
    return jsonify({"ok": True, "llm": llm_result, "report": report, "added_case": _case_to_public_dict(added_case) if added_case else None})


@app.post("/api/upload")
def api_upload():
    if "file" not in request.files:
        return jsonify({"ok": False, "error": "没有收到文件字段 file"}), 400
    file = request.files["file"]
    if not file.filename:
        return jsonify({"ok": False, "error": "文件名为空"}), 400

    original_name = file.filename
    suffix = Path(original_name).suffix.lower()
    safe_stem = Path(secure_filename(original_name)).stem or "upload"
    if suffix not in ALL_UPLOAD_EXTENSIONS:
        return jsonify({"ok": False, "error": "暂不支持该文件类型。支持：jpg/jpeg/png/webp/bmp/tif/tiff/dcm/pdf/txt/md/xlsx/xls/csv/docx。"}), 400

    stored_name = f"{uuid.uuid4().hex}_{safe_stem}{suffix}"
    dest = UPLOAD_DIR / stored_name
    file.save(dest)

    response = {
        "ok": True,
        "filename": original_name,
        "stored_as": stored_name,
        "url": f"/uploads/{stored_name}",
        "size_bytes": dest.stat().st_size,
        "type": "attachment",
        "fields": {},
        "note": "文件已保存。",
    }

    if suffix in CASE_TABLE_EXTENSIONS:
        parsed = parse_case_file(dest)
        candidates = _extract_case_batch_candidates(dest, original_name) if suffix in CASE_TABLE_EXTENSIONS else []
        if candidates:
            batch_id = candidates[0].get("batch_id")
            response.update({
                "type": "candidate_case_batch",
                "batch_id": batch_id,
                "candidate_count": len(candidates),
                "candidate_preview": _json_safe(candidates[:2000]),
                "fields": parsed.get("fields", {}) if parsed.get("ok") else {},
                "parse_ok": True,
                "note": f"已读取病例数据表，共识别 {len(candidates)} 条候选病例。系统已生成候选病例确认卡片，医生可直接选择；也可继续输入年龄、姓名、诊断、病理或住院号等线索后再匹配。",
            })
        elif not parsed.get("ok"):
            response.update({"type": "case_file", "parse_ok": False, "note": parsed.get("error", "病例表格解析失败。"), "fields": {}})
        else:
            response.update({
                "type": "case_file",
                "parse_ok": True,
                "fields": parsed.get("fields", {}),
                "notes": parsed.get("notes", []),
                "field_count": parsed.get("field_count", 0),
                "note": f"已自动识别 {parsed.get('field_count', 0)} 个字段，并纳入当前对话上下文。",
            })
    elif suffix in IMAGE_EXTENSIONS:
        img_kind = "病理/影像图片" if suffix != ".dcm" else "DICOM 影像文件"
        response.update({
            "type": "image",
            "note": f"{img_kind}已保存。若配置了多模态 API，后续问答会把该图片作为当前对话附件传入模型。",
        })
    elif suffix == ".pdf":
        response.update({"type": "pdf", "note": "PDF 已保存。建议同时粘贴关键报告文字以提高检索准确性。"})

    return jsonify(response)


@app.get("/candidate/<batch_id>/<int:row_index>")
def candidate_detail(batch_id, row_index):
    rows = _load_candidate_batch(batch_id)
    candidate = next((x for x in rows if int(x.get("row_index", -1)) == row_index), None)
    if not candidate:
        abort(404)
    return render_template("candidate_detail.html", candidate=candidate)

@app.post("/api/candidate/add")
def api_candidate_add():
    payload = request.get_json(force=True, silent=True) or {}
    candidate = payload.get("candidate") or {}
    if not candidate:
        return jsonify({"ok": False, "error": "未收到候选病例。"}), 400
    source_text = str(candidate.get("free_text") or "")
    duplicate = _find_duplicate_case(candidate, source_text)
    if duplicate:
        return jsonify({"ok": True, "duplicate": True, "case": _case_to_public_dict(duplicate), "message": f"检测到相同或高度相似病例：{duplicate.case_id}，已跳过重复保存。"})
    rec = _add_case_from_fields(candidate, source_text=source_text)
    return jsonify({"ok": True, "case": _case_to_public_dict(rec), "message": f"已将候选病例加入知识库：{rec.case_id}"})


@app.get("/healthz")
def healthz():
    return jsonify({"status": "ok", "records": len(kb.records), "data_path": str(DATA_PATH)})


if __name__ == "__main__":
    port = int(os.getenv("PORT", "5000"))
    debug = os.getenv("FLASK_DEBUG", "0") == "1"
    app.run(host="0.0.0.0", port=port, debug=debug)
