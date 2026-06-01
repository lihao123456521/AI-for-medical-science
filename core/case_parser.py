from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple
import math
import re

import pandas as pd

FIELD_ALIASES: Dict[str, List[str]] = {
    "sex": ["性别", "sex", "gender", "患者性别"],
    "age": ["年龄", "年龄（岁）", "年龄(岁)", "age", "患者年龄"],
    "ls": ["ls", "硬化性苔藓", "lichen sclerosus", "白斑", "有无硬化性苔藓", "硬化性苔藓状态"],
    "tnm": ["tnm", "分期", "临床分期", "病理分期", "tnm分期", "tnm分期/grade分级"],
    "grade": ["grade", "分化", "分化程度", "肿瘤分化程度", "分级"],
    "lymph_node": ["淋巴结", "淋巴转移", "腹股沟淋巴结", "有无淋巴结转移", "n分期"],
    "history": ["病史", "病史要点", "现病史", "既往史", "主诉", "临床病史"],
    "prior_operation": ["既往尿道操作", "既往手术", "尿道扩张", "尿道成形", "尿道重建", "手术史", "操作史", "既往尿道重建手术"],
    "symptoms": ["症状", "体征", "症状/体征", "临床表现", "主诉"],
    "tumor": ["肿瘤", "肿块", "肿物", "占位", "局部描述", "肿瘤情况", "病灶位置", "肿瘤位置"],
    "imaging": ["影像", "影像学", "ct", "mri", "超声", "尿道镜", "内镜", "影像信息", "检查所见"],
    "pathology": ["病理", "活检", "免疫组化", "p63", "ck5/6", "p40", "p16", "ki-67", "病理诊断", "病理信息"],
    "surgery": ["手术", "手术治疗", "术式", "治疗方式", "外科治疗"],
    "other_treatment": ["用药", "药物", "化疗", "放疗", "免疫治疗", "辅助治疗", "其他治疗"],
    "followup": ["随访", "转归", "预后", "复发", "结局"],
    "free_text": ["备注", "补充", "病例摘要", "病例描述", "其他", "完整病例", "讨论"],
}

IMPORTANT_KEYWORDS = {
    "sex", "age", "ls", "history", "symptoms", "tumor", "imaging", "pathology",
    "prior_operation", "tnm", "grade", "lymph_node", "surgery", "other_treatment", "followup", "free_text"
}


def _clean(value: Any) -> str:
    if value is None:
        return ""
    try:
        if isinstance(value, float) and math.isnan(value):
            return ""
    except Exception:
        pass
    text = str(value).strip()
    if text.lower() in {"nan", "none", "null", "nat", "无", "/", "\\"}:
        return ""
    return re.sub(r"\s+", " ", text)


def _norm_key(value: Any) -> str:
    text = _clean(value).lower()
    text = text.replace(" ", "").replace("_", "").replace("-", "").replace("：", "").replace(":", "")
    return text


def _match_field(label: Any) -> str | None:
    key = _norm_key(label)
    if not key:
        return None
    # First pass: exact alias matches. This avoids mapping short labels such as “病理”
    # to broader aliases like “病理分期”.
    for field, aliases in FIELD_ALIASES.items():
        for alias in aliases:
            if key == _norm_key(alias):
                return field
    # Second pass: the alias appears inside a longer column title, e.g. “患者年龄（岁）”.
    for field, aliases in FIELD_ALIASES.items():
        for alias in aliases:
            a = _norm_key(alias)
            if a and a in key:
                return field
    # Third pass: only allow reverse containment for sufficiently specific labels.
    if len(key) >= 3:
        for field, aliases in FIELD_ALIASES.items():
            for alias in aliases:
                a = _norm_key(alias)
                if a and key in a:
                    return field
    return None


def _merge_field(result: Dict[str, str], field: str, value: str) -> None:
    value = _clean(value)
    if not value or field not in IMPORTANT_KEYWORDS:
        return
    # For compact fields, keep the first strong value. For text fields, append distinct chunks.
    if field in {"sex", "age", "ls", "tnm", "grade"}:
        if not result.get(field):
            result[field] = value[:500]
        return
    old = result.get(field, "")
    if old and value in old:
        return
    if old:
        result[field] = (old + "\n" + value)[:4000]
    else:
        result[field] = value[:4000]


def _infer_from_text(text: str, result: Dict[str, str], add_long_text: bool = True) -> None:
    text = _clean(text)
    if not text:
        return
    if not result.get("sex"):
        m = re.search(r"(男|男性|女|女性)", text)
        if m:
            result["sex"] = "男" if "男" in m.group(1) else "女"
    if not result.get("age"):
        m = re.search(r"(\d{1,3})\s*(?:岁|y|yr|year|years|年龄)", text, re.I)
        if m:
            result["age"] = m.group(1)
    lower = text.lower()
    if not result.get("ls") and any(k in lower for k in ["硬化性苔藓", "lichen sclerosus", "ls", "白斑"]):
        result["ls"] = "疑似 LS"
    if not result.get("tnm"):
        m = re.search(r"T\s*[0-4][a-cA-C]?\s*N\s*[0-3xX]?\s*M\s*[01xX]?", text, re.I)
        if m:
            result["tnm"] = re.sub(r"\s+", "", m.group(0).upper())
    if not result.get("grade"):
        m = re.search(r"(?:G|grade)\s*[1-4]|低分化|中分化|高分化", text, re.I)
        if m:
            result["grade"] = m.group(0)

    if not add_long_text:
        return

    # If text is not already assigned through labels, place it into the best long-text bucket.
    if any(k in lower for k in ["mri", "ct", "超声", "尿道镜", "影像", "增强", "淋巴结"]):
        _merge_field(result, "imaging", text)
    elif any(k in lower for k in ["病理", "活检", "鳞癌", "scc", "p63", "ck5/6", "p40", "ki-67", "免疫组化", "异型增生"]):
        _merge_field(result, "pathology", text)
    elif any(k in lower for k in ["排尿", "血尿", "尿潴留", "疼痛", "溃疡", "尿线", "肿物", "肿块"]):
        _merge_field(result, "symptoms", text)
    elif any(k in lower for k in ["尿道扩张", "尿道成形", "尿道重建", "会阴造口", "手术史"]):
        _merge_field(result, "prior_operation", text)
    else:
        _merge_field(result, "free_text", text)


def _parse_key_value_rows(df: pd.DataFrame, result: Dict[str, str]) -> int:
    changed = 0
    if df.empty:
        return changed
    for _, row in df.iterrows():
        values = [_clean(v) for v in row.tolist()]
        values = [v for v in values if v]
        if len(values) < 2:
            continue
        field = _match_field(values[0])
        if field:
            before = result.get(field, "")
            _merge_field(result, field, "；".join(values[1:]))
            changed += int(result.get(field, "") != before)
    return changed


def _parse_header_table(df: pd.DataFrame, result: Dict[str, str]) -> int:
    changed = 0
    if df.empty:
        return changed
    # Drop fully blank rows/columns.
    df = df.dropna(how="all").dropna(axis=1, how="all")
    if df.empty:
        return changed
    headers = [_clean(x) for x in df.iloc[0].tolist()]
    if len(headers) < 2:
        return changed
    # Pick the first non-header-looking row as values.
    for ridx in range(1, min(len(df), 6)):
        row_values = [_clean(x) for x in df.iloc[ridx].tolist()]
        if sum(bool(x) for x in row_values) < 2:
            continue
        for h, v in zip(headers, row_values):
            field = _match_field(h)
            if field and v:
                before = result.get(field, "")
                _merge_field(result, field, v)
                changed += int(result.get(field, "") != before)
        break
    return changed


def parse_excel(path: Path) -> Tuple[Dict[str, str], List[str]]:
    notes: List[str] = []
    result: Dict[str, str] = {}
    suffix = path.suffix.lower()
    if suffix == ".csv":
        sheets = {"CSV": pd.read_csv(path, header=None)}
    else:
        sheets = pd.read_excel(path, sheet_name=None, header=None)
    for sheet_name, df in sheets.items():
        before = dict(result)
        _parse_key_value_rows(df, result)
        _parse_header_table(df, result)
        # Also scan text cells for obvious age/sex/LS/TNM clues.
        text_cells = "\n".join(_clean(x) for x in df.astype(str).values.flatten().tolist() if _clean(x))
        _infer_from_text(text_cells[:8000], result, add_long_text=False)
        if result != before:
            notes.append(f"已解析工作表：{sheet_name}")
    return result, notes


def parse_docx(path: Path) -> Tuple[Dict[str, str], List[str]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx 依赖，请运行 pip install -r requirements.txt 后重试。") from exc

    doc = Document(str(path))
    result: Dict[str, str] = {}
    notes: List[str] = []

    for t_index, table in enumerate(doc.tables, 1):
        rows = [[_clean(cell.text) for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        # Key-value style table: first column label, remaining columns value.
        for row in rows:
            vals = [v for v in row if v]
            if len(vals) >= 2:
                field = _match_field(vals[0])
                if field:
                    _merge_field(result, field, "；".join(vals[1:]))
        # Header style table: first row labels, second row values.
        if len(rows) >= 2:
            headers = rows[0]
            values = rows[1]
            for h, v in zip(headers, values):
                field = _match_field(h)
                if field and v:
                    _merge_field(result, field, v)
        notes.append(f"已扫描 Word 表格 {t_index}")

    para_text = "\n".join(_clean(p.text) for p in doc.paragraphs if _clean(p.text))
    if para_text:
        _infer_from_text(para_text[:8000], result)
        notes.append("已扫描 Word 正文段落")
    return result, notes


def parse_txt(path: Path) -> Tuple[Dict[str, str], List[str]]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    result: Dict[str, str] = {}
    # Parse lines like “年龄：56” or “病理: 鳞癌”。
    for line in text.splitlines():
        if "：" in line:
            k, v = line.split("：", 1)
        elif ":" in line:
            k, v = line.split(":", 1)
        else:
            continue
        field = _match_field(k)
        if field:
            _merge_field(result, field, v)
    _infer_from_text(text[:10000], result)
    return result, ["已解析 TXT 文本"]



def parse_pdf(path: Path) -> Tuple[Dict[str, str], List[str]]:
    try:
        from pypdf import PdfReader
    except Exception:
        try:
            from PyPDF2 import PdfReader
        except Exception as exc:
            raise RuntimeError("PDF 病例解析需要安装 pypdf：pip install pypdf") from exc
    reader = PdfReader(str(path))
    parts: List[str] = []
    for page in reader.pages[:60]:
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        if txt.strip():
            parts.append(txt.strip())
    text = "\n".join(parts)
    result: Dict[str, str] = {}
    _infer_from_text(text[:20000], result)
    return result, [f"已解析 PDF 文本，共 {len(text)} 字符"]

def parse_case_file(path: Path) -> Dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix in {".xlsx", ".xls", ".csv"}:
        fields, notes = parse_excel(path)
    elif suffix == ".docx":
        fields, notes = parse_docx(path)
    elif suffix == ".pdf":
        fields, notes = parse_pdf(path)
    elif suffix in {".txt", ".md"}:
        fields, notes = parse_txt(path)
    elif suffix == ".doc":
        return {
            "ok": False,
            "error": "暂不支持旧版 .doc 二进制格式。请在 Word 中另存为 .docx 后上传。",
            "fields": {},
            "notes": [],
        }
    else:
        return {
            "ok": False,
            "error": "该文件类型不适合作为病例解析。请上传 .xlsx/.xls/.csv/.docx/.pdf/.txt/.md。",
            "fields": {},
            "notes": [],
        }

    # Normalize a few compact fields.
    if "sex" in fields:
        if "男" in fields["sex"]:
            fields["sex"] = "男"
        elif "女" in fields["sex"]:
            fields["sex"] = "女"
    if "age" in fields:
        m = re.search(r"\d{1,3}", fields["age"])
        fields["age"] = m.group(0) if m else fields["age"]

    return {
        "ok": True,
        "fields": fields,
        "notes": notes or ["文件已读取，但未识别到明确字段标签；已尝试从自由文本中抽取。"],
        "field_count": len([v for v in fields.values() if v]),
    }
